import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX

def add_highlighted_heading(doc, text):
    """Add a bold, highlighted heading"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    # Add yellow highlight using shading
    try:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    except:
        # If highlight doesn't work, use shading
        para_format = para.paragraph_format
        para_format.shading.background_color = RGBColor(255, 255, 0)
    return para

def add_section_heading(doc, text):
    """Add a section heading"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    return para

def add_separator(doc, text):
    """Add a separator line"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return para

def main():
    # Read the JSON file
    json_file = "Public/Data/Coriolanus.json"
    print(f"Reading {json_file}...")
    
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Create a new Document
    doc = Document()
    
    # Add title: "Coriolanus" (bold, highlighted)
    add_highlighted_heading(doc, "Coriolanus")
    doc.add_paragraph()  # Add spacing
    
    # Find all ACT and SCENE keys in the data (excluding DRAMATIS PERSONAE)
    scene_keys = [key for key in data.keys() if key.startswith("ACT")]
    # Sort scenes by act and scene number
    def sort_key(key):
        parts = key.split()
        act_num = int(parts[1])
        scene_num = int(parts[3])
        return (act_num, scene_num)
    
    scene_keys.sort(key=sort_key)
    
    print(f"Found {len(scene_keys)} scenes to process")
    
    # Process each scene
    for scene_key in scene_keys:
        # Convert "ACT 1 SCENE 1" to "ACT 1, SCENE 1"
        parts = scene_key.split()
        scene_title = f"{parts[0]} {parts[1]}, {parts[2]} {parts[3]}"
        
        print(f"Processing {scene_title}...")
        
        # Add ACT and SCENE heading
        add_section_heading(doc, scene_title)
        doc.add_paragraph()  # Add spacing
        
        # Add "======PLAY TEXT======" separator
        add_separator(doc, "======PLAY TEXT======")
        doc.add_paragraph()  # Add spacing
        
        # Add all play text lines with line numbers
        scene_data = data[scene_key]
        # Sort by line number
        line_numbers = sorted([int(k) for k in scene_data.keys()])
        
        for line_num in line_numbers:
            line_key = str(line_num)
            if line_key in scene_data:
                play_text = scene_data[line_key].get("play", "")
                if play_text:  # Only add non-empty lines
                    # Add line number and play text
                    para = doc.add_paragraph()
                    run_num = para.add_run(f"{line_num}. ")
                    run_num.font.size = Pt(11)
                    run_text = para.add_run(play_text)
                    run_text.font.size = Pt(11)
        
        doc.add_paragraph()  # Add spacing
        
        # Add "========SCHOLARLY COMMENTARY========" separator
        add_separator(doc, "========SCHOLARLY COMMENTARY========")
        doc.add_paragraph()  # Add spacing
        
        # Leave scholarly commentary section empty (just add a blank paragraph)
        doc.add_paragraph()  # Empty space for scholarly commentary
        
        doc.add_paragraph()  # Add spacing
        doc.add_paragraph()  # Add extra spacing between scenes
    
    # Save the document
    output_file = "Coriolanus_With_Line_Numbers.docx"
    doc.save(output_file)
    print(f"\nDocument saved as: {output_file}")
    print("Done!")

if __name__ == "__main__":
    main()





