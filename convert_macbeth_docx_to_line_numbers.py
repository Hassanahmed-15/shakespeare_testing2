#!/usr/bin/env python3
"""
Convert a Macbeth DOCX (traditional script format) to Macbeth_With_Line_Numbers.docx format.

Source format (from image): ACT 1, Scene 1/2, stage directions italic, speaker names in caps,
sometimes on own line, line numbers (5, 10, 15...) on right margin.

Target format: ACT X, SCENE Y → ========PLAY TEXT======== → 1. line, 2. line, ... → SCHOLARLY COMMENTARY

Usage:
  python convert_macbeth_docx_to_line_numbers.py input.docx
  python convert_macbeth_docx_to_line_numbers.py input.docx -o Macbeth_With_Line_Numbers.docx
"""
import argparse
import os
import re
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_paragraph_info(para):
    """Return (text, is_italic) for a paragraph."""
    text_parts = []
    is_italic = False
    for run in para.runs:
        text_parts.append(run.text)
        if run.italic:
            is_italic = True
    text = "".join(text_parts).strip()
    return text, is_italic


def parse_source_txt(txt_path):
    """
    Parse source TXT (Folger Shakespeare format) into scenes.
    Only SPEAKER is capitalized with colon before dialogue. Skip preamble until ACT 1.
    """
    with open(txt_path, "r", encoding="utf-8", errors="replace") as f:
        raw_lines = [line.rstrip() for line in f]
    scenes = []
    current_act = 1
    current_scene = 1
    current_lines = []
    pending_speaker = None
    started = False
    seen_first_scene = False  # Skip preamble until we see ACT N + Scene N
    stage_pattern = re.compile(
        r"^(Enter |Exit\.?|Exeunt\.?|Alarum|Flourish|Thunder|Lightning\.|They exit\.?|\[.*\]|Sennet\.?)",
        re.I
    )

    for text in raw_lines:
        text = text.strip()
        if not text:
            continue
        # Strip Folger right-margin line numbers (tab + digits) - NOT "ACT 1" or "Scene 1"
        text = re.sub(r"\t\d{1,3}\s*$", "", text).strip()
        if not text:
            continue

        act_match = re.match(r"^ACT\s+([IVXLC]+|\d+)\s*$", text, re.I)
        if act_match:
            started = True
            if seen_first_scene and current_lines:
                scenes.append((current_act, current_scene, current_lines))
                current_lines = []
            current_act = act_match.group(1)
            if str(current_act).isdigit():
                current_act = int(current_act)
            else:
                roman = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5}
                current_act = roman.get(str(current_act).upper(), 1)
            current_scene = 1
            continue

        scene_match = re.match(r"^Scene\s+([IVXLC]+|\d+)\s*$", text, re.I)
        if scene_match:
            seen_first_scene = True
            if current_lines:
                scenes.append((current_act, current_scene, current_lines))
                current_lines = []
            current_scene = scene_match.group(1)
            if str(current_scene).isdigit():
                current_scene = int(current_scene)
            else:
                roman = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6}
                current_scene = roman.get(str(current_scene).upper(), 1)
            continue

        if not seen_first_scene:
            continue

        if stage_pattern.match(text) or (text.startswith("(") and text.endswith(")")):
            line = f"[{text}]" if not text.startswith("[") else text
            if pending_speaker:
                current_lines.append(f"{pending_speaker}: {line}")
                pending_speaker = None
            else:
                current_lines.append(line)
            continue

        words = text.split()
        all_caps = text == text.upper() and len(words) <= 4 and len(text) < 30
        if all_caps and len(words) >= 1:
            pending_speaker = text
            continue

        caps_match = re.match(r"^([A-Z][A-Z\s\'\.\-]{0,30})\s+(.+)$", text)
        if caps_match:
            speaker, rest = caps_match.group(1).strip(), caps_match.group(2).strip()
            stop = {"A", "AN", "THE", "AND", "OR", "BUT", "FOR", "TO", "OF", "IN", "ON", "LIKE"}
            if not any(w in stop for w in speaker.upper().split()):
                current_lines.append(f"{speaker}: {rest}")
                pending_speaker = None
                continue

        if pending_speaker:
            current_lines.append(f"{pending_speaker}: {text}")
            pending_speaker = None
        else:
            current_lines.append(text)

    if current_lines:
        scenes.append((current_act, current_scene, current_lines))
    return scenes


def parse_source_docx(doc_path):
    """
    Parse source DOCX into scenes. Each scene is a list of play lines (strings).
    Returns: list of (act_num, scene_num, lines)
    """
    doc = Document(doc_path)
    scenes = []
    current_act = 1
    current_scene = 1
    current_lines = []
    prev_was_speaker_only = False
    pending_speaker = None

    for para in doc.paragraphs:
        text, is_italic = get_paragraph_info(para)
        if not text:
            continue

        # Strip right-margin line numbers (tab + digits, e.g. "\t10", "\t15")
        text = re.sub(r"\t\d{1,3}\s*$", "", text).strip()

        # Detect ACT heading (e.g. "ACT 1", "ACT I")
        act_match = re.match(r"^ACT\s+([IVXLC]+|\d+)\s*$", text, re.I)
        if act_match:
            started = True
            if current_lines:
                scenes.append((current_act, current_scene, current_lines))
                current_lines = []
            current_act = act_match.group(1)
            if current_act.isdigit():
                current_act = int(current_act)
            else:
                roman = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5}
                current_act = roman.get(str(current_act).upper(), 1)
            current_scene = 1
            continue

        # Detect Scene heading (e.g. "Scene 1", "Scene 2")
        scene_match = re.match(r"^Scene\s+([IVXLC]+|\d+)\s*$", text, re.I)
        if scene_match:
            if current_lines:
                scenes.append((current_act, current_scene, current_lines))
                current_lines = []
            current_scene = scene_match.group(1)
            if str(current_scene).isdigit():
                current_scene = int(current_scene)
            else:
                roman = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6}
                current_scene = roman.get(str(current_scene).upper(), 1)
            continue

        # Stage direction (italic) → wrap in brackets
        if is_italic:
            line = f"[{text}]"
            if pending_speaker:
                current_lines.append(f"{pending_speaker}: {line}")
                pending_speaker = None
            else:
                current_lines.append(line)
            prev_was_speaker_only = False
            continue

        # Check if line is ALL CAPS speaker name only (common names: 1-4 words)
        words = text.split()
        all_caps = text == text.upper() and len(words) <= 4 and all(
            w.replace("'", "").replace(".", "").isalpha() or w in ("FIRST", "SECOND", "THIRD", "ALL")
            for w in words
        )
        if all_caps and len(words) >= 1 and len(text) < 30:
            # Could be speaker on own line - check if next line is dialogue
            pending_speaker = text
            prev_was_speaker_only = True
            continue

        # Speaker + dialogue on same line (e.g. "CAPTAIN Doubtful it stood,")
        caps_match = re.match(r"^([A-Z][A-Z\s\'\.\-]{0,30})\s+(.+)$", text)
        if caps_match:
            speaker = caps_match.group(1).strip()
            rest = caps_match.group(2).strip()
            # Only treat as speaker if it looks like a name (short, no common words)
            stop = {"A", "AN", "THE", "AND", "OR", "BUT", "FOR", "TO", "OF", "IN", "ON", "LIKE"}
            if not any(w in stop for w in speaker.upper().split()):
                line = f"{speaker}: {rest}"
                current_lines.append(line)
                pending_speaker = None
                prev_was_speaker_only = False
                continue

        # Pending speaker from previous line
        if pending_speaker:
            current_lines.append(f"{pending_speaker}: {text}")
            pending_speaker = None
        else:
            current_lines.append(text)
        prev_was_speaker_only = False

    if current_lines:
        scenes.append((current_act, current_scene, current_lines))

    return scenes


def add_highlighted_heading(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    try:
        from docx.enum.text import WD_COLOR_INDEX
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    except Exception:
        pass


def add_section_heading(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def add_separator(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


# Words that are NOT speaker names (verse fragments, common words)
SPEAKER_STOP_WORDS = {
    'A', 'AN', 'THE', 'AND', 'OR', 'BUT', 'FOR', 'TO', 'OF', 'IN', 'ON', 'LIKE',
    'UPON', 'SHOW', 'SHOWED', 'REDOUBLED', 'DOUBLY', 'WHORE', 'REBEL', 'SKINNY',
    'STROKES', 'FOE', 'HIS', 'HER', 'THEIR', 'OUR', 'YOUR', 'MY', 'WHAT', 'WHO',
}


def is_speaker_name(candidate):
    """Return True if candidate looks like a speaker name (e.g. FIRST WITCH, MACBETH)."""
    words = candidate.strip().split()
    if not (1 <= len(words) <= 4) or len(candidate) > 35:
        return False
    if candidate != candidate.upper():
        return False
    return not any(w in SPEAKER_STOP_WORDS for w in words)


def format_line_for_output(line):
    """
    Ensure only SPEAKER is capitalized with colon before dialogue.
    Folger source: only speaker names should be ALL CAPS + ': '
    Dialogue must stay in normal sentence case.
    """
    if not line:
        return line
    # Stage directions - keep as-is
    if line.strip().startswith("["):
        return line
    # Match "SPEAKER: dialogue" - speaker stays caps, ensure dialogue isn't all-caps
    m = re.match(r"^([A-Z][A-Z\s\'\.\-]+):\s*(.*)$", line)
    if m:
        speaker, dialogue = m.group(1).strip(), m.group(2)
        # If dialogue is erroneously all-caps, convert to sentence case
        if dialogue and not dialogue.startswith("["):
            letters = [c for c in dialogue if c.isalpha()]
            if letters and all(c.isupper() for c in letters):
                dialogue = dialogue[0].upper() + dialogue[1:].lower() if len(dialogue) > 1 else dialogue
        return f"{speaker}: {dialogue}".rstrip()
    # No speaker - if entire line is all-caps, convert to sentence case
    if line and not line.startswith("["):
        letters = [c for c in line if c.isalpha()]
        if letters and all(c.isupper() for c in letters):
            return line[0].upper() + line[1:].lower() if len(line) > 1 else line
    return line


def write_target_docx(scenes, output_path, title="Macbeth"):
    doc = Document()
    add_highlighted_heading(doc, title)
    doc.add_paragraph()

    for act_num, scene_num, lines in scenes:
        add_section_heading(doc, f"ACT {act_num}, SCENE {scene_num}")
        doc.add_paragraph()
        add_separator(doc, "======PLAY TEXT======")
        doc.add_paragraph()

        for i, line in enumerate(lines, 1):
            para = doc.add_paragraph()
            run_num = para.add_run(f"{i}. ")
            run_num.font.size = Pt(11)
            formatted = format_line_for_output(line)
            # Bold speaker names (e.g. FIRST WITCH:, MACBETH:) to match target format
            speaker_match = re.match(r"^([A-Z][A-Z '\-\.]+):(.*)$", formatted or "")
            if speaker_match and is_speaker_name(speaker_match.group(1)):
                run_speaker = para.add_run(speaker_match.group(1) + ": ")
                run_speaker.bold = True
                run_speaker.font.size = Pt(11)
                rest = speaker_match.group(2).lstrip()
                if rest:
                    run_rest = para.add_run(rest)
                    run_rest.font.size = Pt(11)
            else:
                run_text = para.add_run(formatted)
                run_text.font.size = Pt(11)

        doc.add_paragraph()
        add_separator(doc, "========SCHOLARLY COMMENTARY========")
        doc.add_paragraph()
        doc.add_paragraph()

    doc.save(output_path)
    print(f"Saved: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Convert Macbeth DOCX to line-numbered format")
    parser.add_argument("input", help="Input DOCX path (your Macbeth script)")
    parser.add_argument("-o", "--output", default="Macbeth_With_Line_Numbers.docx", help="Output DOCX path")
    parser.add_argument("-t", "--title", default="Macbeth", help="Play title")
    args = parser.parse_args()

    script_dir = os.path.dirname(os.path.abspath(__file__))
    input_path = os.path.join(script_dir, args.input) if not os.path.isabs(args.input) else args.input
    output_path = os.path.join(script_dir, args.output) if not os.path.isabs(args.output) else args.output

    if not os.path.isfile(input_path):
        print(f"Input file not found: {input_path}")
        sys.exit(1)

    # For .doc (legacy Word), convert to .txt via textutil
    if input_path.lower().endswith(".doc"):
        import subprocess
        txt_path = input_path[:-4] + ".txt"
        if not os.path.isfile(txt_path) or os.path.getmtime(txt_path) < os.path.getmtime(input_path):
            print(f"Converting .doc to .txt via textutil...")
            subprocess.run(["textutil", "-convert", "txt", input_path, "-output", txt_path], check=True)
        input_path = txt_path

    print(f"Reading {input_path}...")
    if input_path.lower().endswith(".txt"):
        scenes = parse_source_txt(input_path)
    else:
        scenes = parse_source_docx(input_path)
    print(f"Found {len(scenes)} scenes")
    for act, scn, lines in scenes:
        print(f"  ACT {act}, SCENE {scn}: {len(lines)} lines")

    print(f"Writing {output_path}...")
    write_target_docx(scenes, output_path, args.title)
    print("Done!")


if __name__ == "__main__":
    main()
