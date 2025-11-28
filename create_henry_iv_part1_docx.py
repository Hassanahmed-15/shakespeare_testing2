import json
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_highlighted_heading(doc, text):
    """Add a bold, highlighted heading"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    from docx.enum.text import WD_COLOR_INDEX

    try:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    except Exception:
        para.paragraph_format.shading.background_color = RGBColor(255, 255, 0)
    return para


def add_section_heading(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    return para


def add_separator(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return para


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    json_file = os.path.join(script_dir, "Public", "Data", "henry_iv_part1.json")
    print(f"Reading {json_file}...")

    with open(json_file, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    add_highlighted_heading(doc, "Henry IV, Part 1")
    doc.add_paragraph()

    scene_keys = [key for key in data.keys() if key.startswith("ACT")]

    def sort_key(key):
        parts = key.split()
        return (int(parts[1]), int(parts[3]))

    scene_keys.sort(key=sort_key)
    print(f"Found {len(scene_keys)} scenes to process")

    for scene_key in scene_keys:
        parts = scene_key.split()
        scene_title = f"{parts[0]} {parts[1]}, {parts[2]} {parts[3]}"
        print(f"Processing {scene_title}...")

        add_section_heading(doc, scene_title)
        doc.add_paragraph()
        add_separator(doc, "======PLAY TEXT======")
        doc.add_paragraph()

        scene_data = data[scene_key]
        line_numbers = sorted(int(k) for k in scene_data.keys())

        for line_num in line_numbers:
            play_text = scene_data[str(line_num)].get("play", "")
            if not play_text:
                continue
            para = doc.add_paragraph()
            run_num = para.add_run(f"{line_num}. ")
            run_num.font.size = Pt(11)
            run_text = para.add_run(play_text)
            run_text.font.size = Pt(11)

        doc.add_paragraph()
        add_separator(doc, "========SCHOLARLY COMMENTARY========")
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

    output_name = os.environ.get("HENRY_IV_PART1_DOCX", "Henry_IV_Part1_With_Line_Numbers.docx")
    output_file = os.path.join(script_dir, output_name)
    doc.save(output_file)
    print(f"\nDocument saved as: {output_file}")
    print("Done!")


if __name__ == "__main__":
    main()

