#!/usr/bin/env python3
"""
Script to convert DOCX file to JSON format matching the structure of other play JSON files.
Handles all edge cases including DRAMATIS PERSONAE notes and scene notes.
"""

import json
import re
import os
from docx import Document
from typing import Dict, List, Optional, Tuple

def parse_line_number(text: str) -> Optional[Tuple[int, str]]:
    """Extract line number from text like '1: text' or '123: text' or '1. text'"""
    # Try format with colon: "1: text"
    match = re.match(r'^(\d+):\s*(.*)$', text.strip())
    if match:
        return int(match.group(1)), match.group(2)
    
    # Try format with period: "1. text" or "1. word] text"
    match = re.match(r'^(\d+)\.\s*(.*)$', text.strip())
    if match:
        return int(match.group(1)), match.group(2)
    
    return None

def parse_dramatis_note(text: str) -> Optional[Tuple[int, str]]:
    """Parse DRAMATIS PERSONAE note formats: '1: text' or '4. text' or '12. text'"""
    # Format: "1: Dramatis Personae] ..."
    match = re.match(r'^(\d+):\s*(.*)$', text.strip())
    if match:
        return int(match.group(1)), match.group(2)
    
    # Format: "4. Jaques] ..." or "12. Le Beu] ..."
    match = re.match(r'^(\d+)\.\s*(.*)$', text.strip())
    if match:
        return int(match.group(1)), match.group(2)
    
    return None

def is_scene_heading(text: str) -> Optional[Tuple[int, int]]:
    """Check if text is a scene heading like 'ACT 1, SCENE 1' or 'ACT 1 SCENE 1'"""
    # Try format: "ACT 1, SCENE 1"
    match = re.match(r'^ACT\s+(\d+),\s*SCENE\s+(\d+)$', text.strip(), re.IGNORECASE)
    if match:
        return (int(match.group(1)), int(match.group(2)))
    
    # Try format: "ACT 1 SCENE 1" (without comma)
    match = re.match(r'^ACT\s+(\d+)\s+SCENE\s+(\d+)$', text.strip(), re.IGNORECASE)
    if match:
        return (int(match.group(1)), int(match.group(2)))
    
    return None

def is_separator(text: str) -> bool:
    """Check if text is a separator line"""
    text = text.strip()
    return text in ["======PLAY TEXT======", "========SCHOLARLY COMMENTARY========", "====NOTES===="]

def get_paragraph_text(para) -> str:
    """Extract text from a paragraph, handling runs"""
    return para.text.strip()

def extract_dramatis_personae_notes(doc: Document, first_act_para_idx: int) -> Dict[str, List[str]]:
    """Extract DRAMATIS PERSONAE notes from paragraphs before first ACT"""
    dramatis_notes = {}
    
    print("\nExtracting DRAMATIS PERSONAE notes...")
    
    # First, look for "====NOTES====" separator
    notes_start_idx = None
    for i in range(0, first_act_para_idx):
        para = doc.paragraphs[i]
        text = get_paragraph_text(para)
        if "====NOTES====" in text or (is_separator(text) and "NOTES" in text):
            notes_start_idx = i + 1  # Start from next paragraph
            break
    
    # If no separator found, look for note patterns
    if notes_start_idx is None:
        for i in range(0, first_act_para_idx):
            para = doc.paragraphs[i]
            text = get_paragraph_text(para)
            # Look for note patterns - notes contain scholarly references like "]"
            # Character list entries don't have "]"
            if ("Dramatis Personae]" in text or 
                (re.match(r'^\d+[:.]\s+.*[\]\]]', text) and "]" in text) or
                (re.match(r'^\d+\.\s+[A-Z]', text) and any(word in text for word in ["WARBURTON", "JOHNSON", "DYCE", "WRIGHT", "CAPELL", "HALLIWELL"]))):
                notes_start_idx = i
                break
    
    if notes_start_idx is None:
        print("  No DRAMATIS PERSONAE notes found")
        return {}
    
    # Make sure we're past the character list (usually ends around paragraph 11)
    # If notes_start_idx is too early, look for a better starting point
    if notes_start_idx < 10:
        # Check if there are better note indicators later
        for i in range(max(10, notes_start_idx), first_act_para_idx):
            para = doc.paragraphs[i]
            text = get_paragraph_text(para)
            if ("Dramatis Personae]" in text or 
                (re.match(r'^\d+\.\s+[A-Z]', text) and "]" in text)):
                notes_start_idx = i
                break
    
    # Parse notes from notes_start_idx to first_act_para_idx
    current_note_num = None
    current_note_text = None
    
    for i in range(notes_start_idx, first_act_para_idx):
        para = doc.paragraphs[i]
        text = get_paragraph_text(para)
        
        if not text:
            # Empty paragraph - continuation
            if current_note_num is not None and current_note_text:
                current_note_text += " "
            continue
        
        # Check for multiple notes in one paragraph (e.g., "12. Le Beu] ... 7. Rowland de Boys] ...")
        # Look for patterns like "number. word]" which indicate note starts
        note_starts = list(re.finditer(r'\b(\d+)\.\s+([A-Z][^\d]+?)(?=\s+\d+\.|$)', text))
        
        if len(note_starts) > 1:
            # Multiple notes in one paragraph
            # Save current note first
            if current_note_num is not None and current_note_text:
                note_num_str = str(current_note_num)
                if note_num_str not in dramatis_notes:
                    dramatis_notes[note_num_str] = []
                dramatis_notes[note_num_str].append(current_note_text.strip())
                current_note_num = None
                current_note_text = None
            
            # Process each note in the paragraph
            for j, match in enumerate(note_starts):
                note_num = int(match.group(1))
                note_start_pos = match.start()
                # Get text from this note to the next note (or end of paragraph)
                if j + 1 < len(note_starts):
                    note_end_pos = note_starts[j + 1].start()
                    note_text = text[note_start_pos:note_end_pos].strip()
                else:
                    note_text = text[note_start_pos:].strip()
                
                # Parse to get just the note content (after number.)
                note_info = parse_dramatis_note(note_text)
                if note_info:
                    _, note_content = note_info
                    note_num_str = str(note_num)
                    if note_num_str not in dramatis_notes:
                        dramatis_notes[note_num_str] = []
                    dramatis_notes[note_num_str].append(note_content.strip())
            continue
        
        # Single note entry
        note_info = parse_dramatis_note(text)
        if note_info:
            # Save previous note if exists
            if current_note_num is not None and current_note_text:
                note_num_str = str(current_note_num)
                if note_num_str not in dramatis_notes:
                    dramatis_notes[note_num_str] = []
                dramatis_notes[note_num_str].append(current_note_text.strip())
            
            # Start new note
            current_note_num, current_note_text = note_info
        elif current_note_num is not None:
            # Continuation of current note
            if current_note_text:
                current_note_text += " " + text
            else:
                current_note_text = text
    
    # Save last note
    if current_note_num is not None and current_note_text:
        note_num_str = str(current_note_num)
        if note_num_str not in dramatis_notes:
            dramatis_notes[note_num_str] = []
        dramatis_notes[note_num_str].append(current_note_text.strip())
    
    print(f"  Found notes for {len(dramatis_notes)} DRAMATIS PERSONAE entries")
    total_notes = sum(len(notes) for notes in dramatis_notes.values())
    print(f"  Total note entries: {total_notes}")
    return dramatis_notes

def extract_scene_notes(doc: Document, notes_start_idx: int, notes_end_idx: int, scene_key: str) -> Dict[str, List[str]]:
    """Extract notes for a specific scene from a NOTES section"""
    scene_notes = {}
    current_line_num = None
    current_note_text = None
    
    # Start from notes_start_idx (which contains "====NOTES====")
    # If the first note is in the same paragraph, handle it
    start_para = doc.paragraphs[notes_start_idx]
    start_text = get_paragraph_text(start_para)
    
    # Check if NOTES separator and first note are in same paragraph
    if "====NOTES====" in start_text:
        # Extract the note part (everything after "====NOTES====")
        parts = start_text.split("====NOTES====", 1)
        if len(parts) > 1 and parts[1].strip():
            # First note is in the same paragraph
            note_text_after_separator = parts[1].strip()
            note_info = parse_line_number(note_text_after_separator)
            if note_info:
                current_line_num, current_note_text = note_info
            else:
                # No line number, might be continuation or malformed
                pass
    
    # Process remaining paragraphs
    for i in range(notes_start_idx + 1, notes_end_idx):
        para = doc.paragraphs[i]
        text = get_paragraph_text(para)
        
        if not text:
            # Empty paragraph - might be continuation
            if current_line_num is not None and current_note_text:
                current_note_text += " "
            continue
        
        # Check if this is a new note entry (starts with line number)
        note_info = parse_line_number(text)
        if note_info:
            # Save previous note if exists
            if current_line_num is not None and current_note_text:
                line_num_str = str(current_line_num)
                if line_num_str not in scene_notes:
                    scene_notes[line_num_str] = []
                scene_notes[line_num_str].append(current_note_text.strip())
            
            # Start new note
            current_line_num, current_note_text = note_info
        elif current_line_num is not None:
            # Continuation of current note
            if current_note_text:
                current_note_text += " " + text
            else:
                current_note_text = text
        elif is_scene_heading(text):
            # Hit next scene heading, stop processing
            break
    
    # Save last note
    if current_line_num is not None and current_note_text:
        line_num_str = str(current_line_num)
        if line_num_str not in scene_notes:
            scene_notes[line_num_str] = []
        scene_notes[line_num_str].append(current_note_text.strip())
    
    return scene_notes

def convert_docx_to_json(docx_path: str, output_json_path: Optional[str] = None, 
                         existing_json_path: Optional[str] = None) -> Dict:
    """
    Convert DOCX file to JSON format.
    
    Args:
        docx_path: Path to the DOCX file
        output_json_path: Path to save the output JSON (defaults to same name as DOCX)
        existing_json_path: Path to existing JSON file to preserve DRAMATIS PERSONAE
    
    Returns:
        Dictionary with the play data
    """
    print(f"Reading DOCX file: {docx_path}")
    doc = Document(docx_path)
    
    # Find first ACT heading to determine where DRAMATIS PERSONAE ends
    first_act_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = get_paragraph_text(para)
        if is_scene_heading(text):
            first_act_idx = i
            break
    
    if first_act_idx is None:
        print("Error: Could not find first ACT heading")
        return {}
    
    # Extract DRAMATIS PERSONAE entries from DOCX
    print("\nExtracting DRAMATIS PERSONAE entries...")
    dramatis_personae = {}
    
    # Find where NOTES section starts (if any)
    notes_start_idx = None
    for i in range(0, first_act_idx):
        para = doc.paragraphs[i]
        text = get_paragraph_text(para)
        if "====NOTES====" in text or (is_separator(text) and "NOTES" in text):
            notes_start_idx = i
            break
    
    # Extract DRAMATIS PERSONAE entries (before NOTES section)
    end_idx = notes_start_idx if notes_start_idx else first_act_idx
    
    # Start from paragraph 0 to include title
    start_idx = 0
    
    # First pass: add title/header (entry 1)
    for i in range(start_idx, min(2, end_idx)):
        para = doc.paragraphs[i]
        text = get_paragraph_text(para).strip()
        
        if not text or "====NOTES====" in text:
            continue
        
        text_upper = text.upper()
        # Skip if it's a numbered entry (we'll handle those separately)
        if parse_dramatis_note(text):
            continue
        
        # Add as entry 1 (title/header)
        if "1" not in dramatis_personae:
            dramatis_personae["1"] = {
                "play": text,
                "notes": []
            }
    
    # Second pass: process numbered entries
    # Track current entry being built (for multi-paragraph entries)
    current_entry_num = None
    current_entry_text = None
    
    for i in range(start_idx, end_idx):
        para = doc.paragraphs[i]
        text = get_paragraph_text(para)
        
        if not text or "====NOTES====" in text:
            # Save current entry if we have one
            if current_entry_num is not None:
                json_entry_num = str(current_entry_num + 1)  # Offset by 1 for title
                dramatis_personae[json_entry_num] = {
                    "play": current_entry_text.strip(),
                    "notes": []
                }
                current_entry_num = None
                current_entry_text = None
            continue
        
        # Check if this is a numbered entry (format: "1:", "2:", "3:", etc.)
        entry_info = parse_dramatis_note(text)
        if entry_info:
            entry_num, entry_content = entry_info
            
            # Check if this might be a continuation of the previous entry
            # (e.g., if previous was entry_num-1 and this content doesn't start with capital or is very short)
            is_continuation = False
            if current_entry_num is not None and entry_num == current_entry_num + 1:
                # Check if content looks like a continuation (doesn't start with capital, or is very short)
                if entry_content:
                    content_stripped = entry_content.strip()
                    if (not content_stripped[0].isupper() if content_stripped else False) or len(content_stripped) < 20:
                        is_continuation = True
            
            if is_continuation and current_entry_num is not None:
                # This is a continuation - merge with previous entry
                if current_entry_text:
                    current_entry_text += " " + entry_content.strip() if entry_content else ""
                else:
                    current_entry_text = entry_content.strip() if entry_content else ""
            else:
                # Save previous entry if we have one
                if current_entry_num is not None:
                    json_entry_num = str(current_entry_num + 1)  # Offset by 1 for title
                    dramatis_personae[json_entry_num] = {
                        "play": current_entry_text.strip(),
                        "notes": []
                    }
                
                # Start new entry
                current_entry_num = entry_num
                current_entry_text = entry_content if entry_content else ""
        elif current_entry_num is not None:
            # This is a continuation of the current entry (multi-paragraph)
            # Check if it looks like a continuation (not a new numbered entry, not a separator)
            text_upper = text.strip().upper()
            if not is_scene_heading(text) and not is_separator(text) and len(text.strip()) > 0:
                # Append to current entry
                if current_entry_text:
                    current_entry_text += " " + text.strip()
                else:
                    current_entry_text = text.strip()
        else:
            # Not a numbered entry and not a continuation
            # Check if it's an unnumbered character entry
            text_upper = text.strip().upper()
            
            # Skip title/header (already added) and separators
            if (text_upper == "DRAMATIS PERSONAE" or 
                is_scene_heading(text) or 
                is_separator(text) or
                i < 1):  # Skip first paragraph (title/header)
                continue
            
            # Unnumbered character entry - use sequential numbering
            max_num = 0
            for key in dramatis_personae.keys():
                if key.isdigit():
                    max_num = max(max_num, int(key))
            
            if text and len(text.strip()) > 3:
                entry_num_str = str(max_num + 1)
                dramatis_personae[entry_num_str] = {
                    "play": text.strip(),
                    "notes": []
                }
    
    # Save final entry if we have one
    if current_entry_num is not None:
        json_entry_num = str(current_entry_num + 1)  # Offset by 1 for title
        dramatis_personae[json_entry_num] = {
            "play": current_entry_text.strip(),
            "notes": []
        }
    
    print(f"  Extracted {len(dramatis_personae)} DRAMATIS PERSONAE entries")
    
    # Always use extracted DRAMATIS PERSONAE from DOCX (it's the source of truth)
    # Only use existing JSON for notes matching if needed
    print(f"  Using extracted DRAMATIS PERSONAE from DOCX ({len(dramatis_personae)} entries)")
    
    # Extract DRAMATIS PERSONAE notes
    dramatis_notes = extract_dramatis_personae_notes(doc, first_act_idx)
    
    # Match notes to DRAMATIS PERSONAE entries by character name mentioned in note
    # Notes reference character numbers that may not match JSON numbering
    if dramatis_personae:
        # Build a map of character names to entry numbers
        char_name_to_entry = {}
        for entry_num, entry_data in dramatis_personae.items():
            play_text = entry_data.get("play", "")
            # Extract character name (usually first word or words before comma)
            if play_text:
                # Get the main character name
                name_parts = play_text.split(",")[0].strip().upper()
                # Store full name
                char_name_to_entry[name_parts] = entry_num
                # Also try without common prefixes
                for prefix in ["THE ", "A ", "SIR ", "LORD ", "LADY "]:
                    if name_parts.startswith(prefix):
                        char_name_to_entry[name_parts[len(prefix):]] = entry_num
                # Also add individual significant words (skip common words)
                for word in name_parts.split():
                    if len(word) > 2 and word not in ["THE", "A", "SIR", "LORD", "LADY", "AND", "OF", "TO"]:
                        char_name_to_entry[word] = entry_num
        
        # Match notes to entries by character name mentioned in note
        # Process all notes and match by name first, never create placeholder entries
        for note_num, notes_list in dramatis_notes.items():
            matched = False
            best_match_entry = None
            best_match_score = 0
            
            for note_text in notes_list:
                # Look for character name in note (usually before "]")
                if "]" in note_text:
                    # Extract the part before "]"
                    name_part = note_text.split("]")[0].strip()
                    # Remove leading numbers and periods (e.g., "4. Jaques]" -> "Jaques")
                    name_part = re.sub(r'^\d+\.\s*', '', name_part).strip().upper()
                    # Remove trailing period if present
                    name_part = name_part.rstrip('.')
                    
                    # Try to match character name - prioritize exact matches
                    for char_name, entry_num in char_name_to_entry.items():
                        # Normalize both for comparison
                        char_name_clean = char_name.rstrip('.')
                        name_part_clean = name_part
                        
                        # Exact match gets highest score
                        if name_part_clean == char_name_clean:
                            best_match_entry = entry_num
                            best_match_score = 100
                            matched = True
                            break
                        # Partial match - check if name_part contains char_name or vice versa
                        elif len(name_part_clean) > 2 and len(char_name_clean) > 2:
                            # Check if one contains the other (more flexible matching)
                            if char_name_clean in name_part_clean:
                                score = len(char_name_clean)
                                if score > best_match_score:
                                    best_match_entry = entry_num
                                    best_match_score = score
                                    matched = True
                            elif name_part_clean in char_name_clean:
                                score = len(name_part_clean)
                                if score > best_match_score:
                                    best_match_entry = entry_num
                                    best_match_score = score
                                    matched = True
                            # Also try matching key words (e.g., "Cardinal" should match "Cardinal (Thomas Bourchier)")
                            elif len(name_part_clean.split()) == 1 and len(char_name_clean.split()) > 1:
                                # Single word in note, check if it's the first significant word in character name
                                first_word = char_name_clean.split()[0]
                                if name_part_clean == first_word:
                                    score = len(name_part_clean)
                                    if score > best_match_score:
                                        best_match_entry = entry_num
                                        best_match_score = score
                                        matched = True
                    
                    if best_match_score >= 100:  # Found exact match, stop searching
                        break
            
            # Apply notes to best match (only if we found a match)
            if matched and best_match_entry and best_match_entry in dramatis_personae:
                if "notes" not in dramatis_personae[best_match_entry]:
                    dramatis_personae[best_match_entry]["notes"] = []
                dramatis_personae[best_match_entry]["notes"].extend(notes_list)
            # DO NOT create placeholder entries - if no match found, skip the note
            # This prevents overwriting real character entries with placeholders
    else:
        # Create placeholder DRAMATIS PERSONAE if not found
        play_name = os.path.basename(docx_path).replace('.docx', '').replace('_', ' ').title()
        dramatis_personae = {
            "1": {
                "play": f"DRAMATIS PERSONAE\n{play_name.upper()}",
                "notes": dramatis_notes.get("1", [])
            }
        }
        for note_num, notes_list in dramatis_notes.items():
            if note_num != "1":
                dramatis_personae[note_num] = {
                    "play": f"DRAMATIS PERSONAE entry {note_num}",
                    "notes": notes_list
                }
    
    # Find all NOTES sections
    notes_sections = []
    for i, para in enumerate(doc.paragraphs):
        text = get_paragraph_text(para)
        if is_separator(text) and ("NOTES" in text or "COMMENTARY" in text):
            notes_sections.append(i)
    
    print(f"\nFound {len(notes_sections)} NOTES sections")
    
    # Parse the document - first pass: extract play text
    all_scenes = {}
    current_scene_key = None
    current_scene_lines = {}
    in_play_text_section = False
    in_commentary_section = False
    skip_dramatis = True  # Skip DRAMATIS PERSONAE section at the beginning
    
    print("\nParsing document (play text)...")
    
    for para in doc.paragraphs:
        text = get_paragraph_text(para)
        
        if not text:
            continue
        
        # Check for scene heading
        scene_info = is_scene_heading(text)
        if scene_info:
            # Save previous scene if exists
            if current_scene_key and current_scene_lines:
                all_scenes[current_scene_key] = current_scene_lines
                print(f"  Completed {current_scene_key}: {len(current_scene_lines)} lines")
            
            # Start new scene
            act_num, scene_num = scene_info
            current_scene_key = f"ACT {act_num} SCENE {scene_num}"
            current_scene_lines = {}
            in_play_text_section = True  # After scene heading, we're in play text
            in_commentary_section = False
            skip_dramatis = False  # We've passed DRAMATIS PERSONAE
            print(f"  Found {current_scene_key}")
            continue
        
        # Skip DRAMATIS PERSONAE section (before first scene)
        if skip_dramatis:
            continue
        
        # Check for separators - NOTES separator might be combined with first note
        is_notes_separator = is_separator(text) or "====NOTES====" in text
        if is_notes_separator or is_separator(text):
            if "PLAY TEXT" in text:
                in_play_text_section = True
                in_commentary_section = False
            elif "COMMENTARY" in text or "NOTES" in text or "====NOTES====" in text:
                in_play_text_section = False
                in_commentary_section = True
            continue
        
        # Skip commentary section for now (we'll process it separately)
        if in_commentary_section:
            continue
        
        # Process lines in play text section (or if no separator, assume we're in play text after scene heading)
        # Skip if this paragraph contains NOTES separator (even if combined with note text)
        if "====NOTES====" in text:
            in_play_text_section = False
            in_commentary_section = True
            continue
        
        # CRITICAL: Check if this looks like a note entry BEFORE processing as play text
        # Notes have patterns like "10. word] Scholar name:" or "word] Scholar:"
        # This check must happen BEFORE we process it as play text
        scholars_check = ["WARBURTON", "JOHNSON", "DYCE", "WRIGHT", "CAPELL", "MOBERLY", "THEOBALD", "COLERIDGE", "STEEVENS", "HALLIWELL", "CALDECOTT", "WALKER", "KNIGHT", "NARES", "ABBOTT", "STEEVENS", "MOBERLY", "THEOBALD"]
        
        # Quick check: if text contains ] and scholar names, it's likely a note
        if "]" in text and any(scholar in text.upper() for scholar in scholars_check):
            # Parse to check line number format
            line_info = parse_line_number(text)
            if line_info:
                _, line_text = line_info
                # Check if this is clearly a note (has ] followed by scholar name)
                if "]" in line_text:
                    parts = line_text.split("]", 1)
                    if len(parts) > 1:
                        after_bracket = parts[1].strip()
                        # Check if it starts with a scholar name or has scholar name pattern
                        if any(scholar in after_bracket.upper()[:50] for scholar in scholars_check):
                            # This is definitely a note - skip it completely
                            # Set commentary section flag to prevent future processing
                            in_commentary_section = True
                            continue
            
        if (in_play_text_section or current_scene_key) and not in_commentary_section:
            line_info = parse_line_number(text)
            if line_info:
                line_num, line_text = line_info
                
                # Clean line_text - remove any note-like content (contains scholarly references)
                # Notes typically contain patterns like "word]" followed by scholar names
                # Check if line_text contains note patterns
                scholars = ["WARBURTON", "JOHNSON", "DYCE", "WRIGHT", "CAPELL", "MOBERLY", "THEOBALD", "COLERIDGE", "STEEVENS", "HALLIWELL", "CALDECOTT", "WALKER", "KNIGHT", "STEEVENS", "NARES", "ABBOTT"]
                
                if "]" in line_text:
                    # Check if there's a note pattern - look for word] followed by scholar name
                    # Split by ] and check each part
                    parts = line_text.split("]")
                    clean_parts = []
                    found_note_start = False
                    
                    for i, part in enumerate(parts):
                        if i == 0:
                            # First part - check if it ends with a word that looks like a note marker
                            # Notes often have patterns like "word]" where word is the subject
                            # But play text might also have "]" in stage directions like "[Enter]"
                            # If the part after ] contains scholar names, this is a note
                            if len(parts) > 1:
                                after_bracket = parts[1].strip()
                                if any(scholar in after_bracket.upper() for scholar in scholars):
                                    # This is note content - keep only text before the note starts
                                    # Find where the note starts (usually a word before ])
                                    # For now, just keep the part before ]
                                    clean_parts.append(part.strip())
                                    found_note_start = True
                                    break
                                else:
                                    clean_parts.append(part)
                            else:
                                clean_parts.append(part)
                        else:
                            # Subsequent parts - check if this contains scholar names
                            if any(scholar in part.upper() for scholar in scholars):
                                # This is note content, stop here
                                found_note_start = True
                                break
                            elif not found_note_start:
                                clean_parts.append("]" + part)
                    
                    if found_note_start:
                        # Reconstruct clean text (only parts before note)
                        line_text = "]".join(clean_parts).strip()
                        # Remove trailing ] if present
                        if line_text.endswith("]"):
                            line_text = line_text[:-1].strip()
                    elif any(scholar in line_text.upper() for scholar in scholars) and "]" in line_text:
                        # Has scholar name and ], likely a note - extract only before ]
                        parts = line_text.split("]", 1)
                        if len(parts) > 1 and any(scholar in parts[1].upper() for scholar in scholars):
                            line_text = parts[0].strip()
                            if line_text.endswith("]"):
                                line_text = line_text[:-1].strip()
                    
                if line_text:  # Only add non-empty lines
                    line_num_str = str(line_num)
                    
                    # Check if this looks like note content (has ] and scholar names)
                    is_note_content = False
                    if "]" in line_text:
                        # Check if it has note pattern
                        parts = line_text.split("]", 1)
                        if len(parts) > 1:
                            after_bracket = parts[1].strip()
                            if any(scholar in after_bracket.upper() for scholar in scholars):
                                is_note_content = True
                    
                    # Skip if this is note content - notes will be added separately
                    if is_note_content:
                        continue
                    
                    if line_num_str not in current_scene_lines:
                        current_scene_lines[line_num_str] = {
                            "play": line_text,
                            "notes": []  # Initialize notes array
                        }
                    else:
                        # Line already exists, might be duplicate or continuation
                        # Only append if it's clearly play text continuation (no note patterns)
                        if not is_note_content:
                            # Check if appending would create note-like content
                            combined = current_scene_lines[line_num_str]["play"] + " " + line_text
                            if not ("]" in combined and any(scholar in combined.upper() for scholar in scholars)):
                                current_scene_lines[line_num_str]["play"] += " " + line_text
            elif text and not is_separator(text) and "====NOTES====" not in text and current_scene_key:
                # Line without number - might be continuation or stage direction
                # Add it with the next available line number
                if current_scene_lines:
                    max_line = max(int(k) for k in current_scene_lines.keys())
                    next_line = str(max_line + 1)
                else:
                    next_line = "1"
                if next_line not in current_scene_lines:
                    current_scene_lines[next_line] = {
                        "play": text,
                        "notes": []  # Initialize notes array
                    }
                else:
                    current_scene_lines[next_line]["play"] += " " + text
    
    # Save last scene
    if current_scene_key and current_scene_lines:
        all_scenes[current_scene_key] = current_scene_lines
        print(f"  Completed {current_scene_key}: {len(current_scene_lines)} lines")
    
    # Second pass: extract notes from all NOTES sections and match them to scenes
    print("\nParsing notes sections...")
    
    # First, build a map of scene positions
    scene_positions = {}  # {scene_key: paragraph_index}
    for i, para in enumerate(doc.paragraphs):
        text = get_paragraph_text(para)
        scene_info = is_scene_heading(text)
        if scene_info:
            act_num, scene_num = scene_info
            scene_key = f"ACT {act_num} SCENE {scene_num}"
            scene_positions[scene_key] = i
    
    # Map NOTES sections to scenes (each NOTES section belongs to the most recent scene before it)
    notes_to_scenes = {}  # {scene_key: [(notes_start_idx, notes_end_idx), ...]}
    
    for i, para in enumerate(doc.paragraphs):
        text = get_paragraph_text(para)
        
        # Check if this is a NOTES separator
        # Check for NOTES separator - it might be standalone or combined with first note
        is_notes_separator = False
        if is_separator(text) and ("NOTES" in text or "COMMENTARY" in text):
            is_notes_separator = True
        elif "====NOTES====" in text:
            # NOTES separator might be in same paragraph as first note
            is_notes_separator = True
        
        if is_notes_separator:
            # Find which scene this NOTES section belongs to
            # It belongs to the most recent scene before this NOTES section
            matching_scene_key = None
            matching_scene_pos = -1
            
            for scene_key, scene_pos in scene_positions.items():
                if scene_pos < i and scene_pos > matching_scene_pos:
                    matching_scene_pos = scene_pos
                    matching_scene_key = scene_key
            
            if matching_scene_key:
                # Find the end of this NOTES section (next scene or next NOTES section)
                end_idx = len(doc.paragraphs)
                for j in range(i + 1, len(doc.paragraphs)):
                    next_text = get_paragraph_text(doc.paragraphs[j])
                    if is_scene_heading(next_text):
                        end_idx = j
                        break
                    # Check for next NOTES section (but not if it's in the same paragraph)
                    if j > i and is_separator(next_text) and "NOTES" in next_text:
                        end_idx = j
                        break
                    elif j > i and "====NOTES====" in next_text and next_text.strip().startswith("====NOTES===="):
                        # Only if it's a new NOTES section, not continuation
                        if not any(char.isdigit() for char in next_text[:50]):  # No line numbers = new section
                            end_idx = j
                            break
                
                if matching_scene_key not in notes_to_scenes:
                    notes_to_scenes[matching_scene_key] = []
                notes_to_scenes[matching_scene_key].append((i, end_idx))
    
    # Extract notes for each scene
    for scene_key, notes_ranges in notes_to_scenes.items():
        if scene_key not in all_scenes:
            print(f"  Warning: Scene {scene_key} not found in all_scenes, skipping notes")
            continue
        
        scene_notes_combined = {}
        for notes_start, notes_end in notes_ranges:
            scene_notes = extract_scene_notes(doc, notes_start, notes_end, scene_key)
            # Merge notes (multiple notes per line number)
            for line_num, notes_list in scene_notes.items():
                if line_num not in scene_notes_combined:
                    scene_notes_combined[line_num] = []
                scene_notes_combined[line_num].extend(notes_list)
        
        # Apply notes to scene lines AND clean play text
        scholars_clean = ["WARBURTON", "JOHNSON", "DYCE", "WRIGHT", "CAPELL", "MOBERLY", "THEOBALD", "COLERIDGE", "STEEVENS", "HALLIWELL", "CALDECOTT", "WALKER", "KNIGHT", "NARES", "ABBOTT"]
        for line_num, notes_list in scene_notes_combined.items():
            if line_num in all_scenes[scene_key]:
                all_scenes[scene_key][line_num]["notes"] = notes_list
                
                # Clean play text - remove any note content that might have slipped through
                play_text = all_scenes[scene_key][line_num].get("play", "")
                if "]" in play_text:
                    # Check if play text contains note patterns
                    parts = play_text.split("]", 1)
                    if len(parts) > 1:
                        after_bracket = parts[1].strip()
                        if any(scholar in after_bracket.upper() for scholar in scholars_clean):
                            # This contains note content - remove it
                            clean_play_text = parts[0].strip()
                            # Remove trailing ] if present
                            if clean_play_text.endswith("]"):
                                clean_play_text = clean_play_text[:-1].strip()
                            all_scenes[scene_key][line_num]["play"] = clean_play_text
        
        # Also clean all play text entries in this scene (even if they don't have notes)
        for line_num, line_data in all_scenes[scene_key].items():
            play_text = line_data.get("play", "")
            if "]" in play_text and any(scholar in play_text.upper() for scholar in scholars_clean):
                # Check if this is note content
                parts = play_text.split("]", 1)
                if len(parts) > 1:
                    after_bracket = parts[1].strip()
                    if any(scholar in after_bracket.upper()[:100] for scholar in scholars_clean):
                        # Remove note content
                        clean_play_text = parts[0].strip()
                        if clean_play_text.endswith("]"):
                            clean_play_text = clean_play_text[:-1].strip()
                        all_scenes[scene_key][line_num]["play"] = clean_play_text
        
        total_notes = sum(len(notes) for notes in scene_notes_combined.values())
        if total_notes > 0:
            print(f"  Added notes to {scene_key}: {total_notes} note entries across {len(scene_notes_combined)} lines")
        else:
            print(f"  Warning: No notes found for {scene_key}")
    
    # Final cleaning pass: Remove any note content from play text
    print("\nCleaning play text (removing any note content)...")
    scholars_final = [
        "WARBURTON", "JOHNSON", "DYCE", "WRIGHT", "CAPELL", "MOBERLY", "THEOBALD", 
        "COLERIDGE", "STEEVENS", "HALLIWELL", "CALDECOTT", "WALKER", "KNIGHT", 
        "NARES", "ABBOTT", "SINGER", "RUSHTON", "CAMPBELL", "HEARD", "SEYMOUR"
    ]
    cleaned_count = 0
    
    for scene_key, scene_data in all_scenes.items():
        for line_num, line_data in scene_data.items():
            play_text = line_data.get("play", "")
            if not play_text:
                continue
            
            # Check if play text contains note patterns
            # Look for: word] followed by explanatory text (scholar names, "Ed.", "A technical", etc.)
            if "]" in play_text:
                # Split by ] and check
                parts = play_text.split("]", 1)  # Split only on first ]
                if len(parts) == 2:
                    before_bracket = parts[0].strip()
                    after_bracket = parts[1].strip()
                    
                    if not after_bracket:
                        continue
                    
                    # Check if after bracket contains note indicators:
                    after_upper = after_bracket.upper()
                    is_note = False
                    
                    # Check for scholar names (check in first 200 chars to catch names)
                    for scholar in scholars_final:
                        if scholar in after_upper[:200]:
                            is_note = True
                            break
                    
                    # Check for editor markers
                    if not is_note:
                        if "—ED." in after_upper or " ED." in after_upper or after_upper.startswith("ED."):
                            is_note = True
                    
                    # Check for note patterns (expanded list)
                    if not is_note:
                        note_patterns = [
                            "A TECHNICAL", "ACCORDING TO", "AS GIVEN BY", "AS IN", 
                            "USED WHERE", "SIGNIFYING", "MEANING", "REFERS TO",
                            "THAT IS", "I.E.", "VIZ.", "SEE", "CF.", "COMPARE",
                            "SUGGESTS", "SCANS", "IS", "ARE", "WAS", "WERE",
                            "SEEMS TO HAVE", "DISTINGUISHED", "DEMANDED", "AS FAR AS",
                            "CRIT.", "LIT.", "ILLUST.", "NICHOLS"
                        ]
                        for pattern in note_patterns:
                            if pattern in after_upper[:150]:  # Check first 150 chars
                                is_note = True
                                break
                    
                    # Check if it's a stage direction (should be preserved)
                    # Stage directions are usually short, contain words like "Aside", "To", "Enter", "Exit"
                    is_stage_direction = False
                    stage_direction_words = ["ASIDE", "TO", "ENTER", "EXIT", "EXEUNT", "WITHIN", "ABOVE", "BELOW"]
                    if len(after_bracket) < 50:  # Stage directions are usually short
                        if any(word in after_upper[:50] for word in stage_direction_words):
                            is_stage_direction = True
                    
                    # More aggressive check: if after bracket is longer than 20 chars and doesn't look like dialogue/stage direction
                    if not is_note and not is_stage_direction and len(after_bracket) > 20:
                        # Check if it doesn't look like dialogue (dialogue usually has quotes, colons, exclamations)
                        first_50 = after_bracket[:50]
                        has_dialogue_markers = any(char in first_50 for char in ['"', "'", ':', '!', '?'])
                        
                        # If it's long text without dialogue markers, it's likely a note
                        if not has_dialogue_markers:
                            # Check for common note words/phrases
                            explanatory_indicators = [
                                "WHICH", "THAT", "SEEMS", "HAVE", "HAS", "HAD",
                                "DISTINGUISHED", "DEMANDED", "REQUIRED", "MEANS",
                                "REFERS", "SIGNIFIES", "INDICATES", "SUGGESTS"
                            ]
                            if any(word in after_upper[:150] for word in explanatory_indicators):
                                is_note = True
                    
                    # Final check: if text after ] starts with a capital letter followed by lowercase (like "According", "That is")
                    # and is longer than 15 chars, it's likely a note (but not if it's a stage direction)
                    if not is_note and not is_stage_direction and len(after_bracket) > 15:
                        first_word = after_bracket.strip().split()[0] if after_bracket.strip() else ""
                        if len(first_word) > 3 and first_word[0].isupper() and first_word[1:].islower():
                            # Check if it's a common note starter (but not stage direction words)
                            note_starters = ["According", "That", "This", "These", "Those", "As", "Used", "Signifying", "Meaning"]
                            if first_word not in ["Aside", "To", "Enter", "Exit"] and any(starter in first_word for starter in note_starters):
                                is_note = True
                    
                    if is_note:
                        # This is note content - keep only the part before ]
                        clean_play_text = before_bracket
                        # Remove trailing ] if present
                        if clean_play_text.endswith("]"):
                            clean_play_text = clean_play_text[:-1].strip()
                        
                        # Also remove trailing comma if the note was attached
                        if clean_play_text.endswith(","):
                            clean_play_text = clean_play_text[:-1].strip()
                        
                        if clean_play_text != play_text:
                            line_data["play"] = clean_play_text
                            cleaned_count += 1
    
    if cleaned_count > 0:
        print(f"  Cleaned note content from {cleaned_count} play text entries")
    else:
        print(f"  No note content found in play text (all clean)")
    
    # Combine DRAMATIS PERSONAE and scenes
    output_data = {
        "DRAMATIS PERSONAE": dramatis_personae
    }
    output_data.update(all_scenes)
    
    # Determine output path
    if not output_json_path:
        base_name = os.path.basename(docx_path).replace('.docx', '').replace('.doc', '')
        # Convert to snake_case for filename
        base_name = base_name.lower().replace(' ', '_').replace('-', '_')
        script_dir = os.path.dirname(os.path.abspath(__file__))
        output_json_path = os.path.join(script_dir, "Public", "Data", f"{base_name}.json")
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(output_json_path), exist_ok=True)
    
    # Write to file
    print(f"\nWriting JSON to: {output_json_path}")
    with open(output_json_path, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, indent=2, ensure_ascii=False)
    
    print(f"\nConversion complete!")
    print(f"  Total scenes: {len(all_scenes)}")
    total_lines = sum(len(scene_data) for scene_data in all_scenes.values())
    print(f"  Total lines: {total_lines}")
    total_notes = sum(
        sum(len(line_data.get("notes", [])) for line_data in scene_data.values())
        for scene_data in all_scenes.values()
    )
    print(f"  Total notes: {total_notes}")
    dramatis_notes_count = sum(
        len(entry.get("notes", [])) for entry in dramatis_personae.values()
    )
    print(f"  DRAMATIS PERSONAE notes: {dramatis_notes_count}")
    print(f"  Output file: {output_json_path}")
    
    return output_data

def main():
    import sys
    
    # Default input file
    default_input = "Public/Data/As_You_Like_It_Clean.docx"
    default_existing = "Public/Data/as_you_like_it.json"
    
    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
    else:
        docx_path = default_input
    
    # Check if file exists
    if not os.path.exists(docx_path):
        print(f"Error: File not found: {docx_path}")
        sys.exit(1)
    
    # Try to use existing JSON for DRAMATIS PERSONAE
    existing_json = None
    if docx_path == default_input:
        existing_json = default_existing
    elif len(sys.argv) > 2:
        existing_json = sys.argv[2]
    
    convert_docx_to_json(docx_path, existing_json_path=existing_json)

if __name__ == "__main__":
    main()
