#!/usr/bin/env python3
"""
Replace Sections III and IV in the paper with expanded methodology.
Preserves all formatting, styles, and surrounding content.
"""

from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import re

DOCX_PATH = "Public/Data/Full Fathom Five Hassan Edition.docx"
SECTION_FONT = Emu(228600)   # 18pt for section headers (III.)
SUBSECTION_FONT = Emu(171450) # 13.5pt for subsection headers (A.)
FONT_NAME = "Times New Roman"

# ── helpers ──────────────────────────────────────────────────────────────

def make_run(parent_p, text, bold=None, italic=None, font_size=None):
    """Append a formatted run to a paragraph element."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if italic:
        i = OxmlElement('w:i')
        rPr.append(i)

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    rPr.append(rFonts)

    if font_size is not None:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size.pt * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
        rPr.append(szCs)

    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    parent_p.append(r)
    return r


def make_paragraph():
    """Create a bare <w:p> element."""
    return OxmlElement('w:p')


def section_header(text):
    """Create a bold 18pt section header paragraph."""
    p = make_paragraph()
    make_run(p, text, bold=True, font_size=Pt(18))
    return p


def subsection_header(text):
    """Create a bold 13.5pt subsection header paragraph."""
    p = make_paragraph()
    make_run(p, text, bold=True, font_size=Pt(13.5))
    return p


def parse_rich_text(markup):
    """
    Parse simple markup into a list of (text, bold, italic) tuples.
    Supported tags: <b>...</b>, <i>...</i>, <bi>...</bi>
    Everything else is plain text.
    """
    runs = []
    pattern = re.compile(r'<(b|i|bi)>(.*?)</\1>', re.DOTALL)
    pos = 0
    for m in pattern.finditer(markup):
        if m.start() > pos:
            runs.append((markup[pos:m.start()], False, False))
        tag = m.group(1)
        txt = m.group(2)
        bold = tag in ('b', 'bi')
        italic = tag in ('i', 'bi')
        runs.append((txt, bold, italic))
        pos = m.end()
    if pos < len(markup):
        runs.append((markup[pos:], False, False))
    return runs


def body_para(markup):
    """Create a body paragraph from markup with <b>, <i>, <bi> tags."""
    p = make_paragraph()
    for text, bold, italic in parse_rich_text(markup):
        make_run(p, text,
                 bold=bold if bold else None,
                 italic=italic if italic else None)
    return p


# ── content definition ───────────────────────────────────────────────────

def build_new_methodology():
    """Return a list of OxmlElement <w:p> paragraphs for the new methodology."""
    paras = []

    # ── SECTION HEADER ──
    paras.append(section_header("III. Methodology and System Development"))

    paras.append(body_para(
        "This section presents a comprehensive account of the technical methodology employed to transform the New Variorum Shakespeare from physical and digitized volumes into an interactive, AI-augmented scholarly platform. The process encompassed corpus acquisition and assessment, multi-stage optical character recognition with large-language-model validation, structured data modeling, editorial alignment across historical and contemporary editions, systematic bibliographic resolution, web platform engineering, serverless AI integration, retrieval-augmented generation for historical commentary, multimedia discovery, and iterative quality assurance. Each subsection details the specific tools, architectural decisions, data structures, and computational strategies used to achieve the project\u2019s objectives."
    ))

    # ── A ──
    paras.append(subsection_header("A. Corpus Acquisition and Digital Source Assessment"))

    paras.append(body_para(
        "The foundation of the project required acquiring high-fidelity digital surrogates of all available pre-1956 New Variorum Shakespeare editions. We conducted a systematic survey of Internet Archive and HathiTrust Digital Library repositories to identify digitized copies of each published volume. Of the 23 dramatic editions published between 1871 (Furness\u2019s <i>Romeo and Juliet</i>) and 1955 (Black\u2019s <i>Richard II</i>), 20 were available as complete, page-level digital scans of sufficient resolution for optical character recognition processing. The three remaining volumes\u2014<i>Henry IV, Part 2</i> (Shaaber, 1940), <i>Troilus and Cressida</i> (Hillebrand and Baldwin, 1953), and <i>Richard II</i> (Black, 1955)\u2014are currently being digitized through the same pipeline and will be incorporated upon completion."
    ))

    paras.append(body_para(
        "Source assessment involved evaluating each digitized volume for scan quality, page completeness, and typographic legibility. The Variorum editions present atypical digitization challenges distinct from standard literary texts: their page layouts intermingle Shakespeare\u2019s dramatic text (set in larger type) with densely packed commentary in smaller type, marginal line numbers, footnote apparatus with abbreviated citations, and occasionally multiple columns of commentary per page. Several volumes exhibited degraded scan quality in footnote regions where microscopic typography and thin ink produced ambiguous character forms. These regions required particular attention during the OCR validation stage described below."
    ))

    paras.append(body_para(
        "Each volume was catalogued with metadata including editor, publication year, publisher, approximate page count, and an assessment of scan quality categorized as high, moderate, or low. This metadata informed prioritization decisions during the OCR pipeline, with lower-quality scans allocated additional validation passes."
    ))

    # ── B ──
    paras.append(subsection_header("B. Multi-Stage OCR Pipeline with Large-Language-Model Validation"))

    paras.append(body_para(
        "The conversion of digitized page images into machine-readable text required a multi-stage pipeline engineered to handle the typographic complexity of historical scholarly editions. Standard OCR engines, while effective on clean modern typography, exhibited elevated error rates when applied to the Variorum\u2019s dense footnote apparatus, abbreviated citations, mixed typefaces (roman, italic, blackletter), and historical printing conventions including ligatures, long-s forms, and inconsistent spacing."
    ))

    paras.append(body_para(
        "<b>Stage 1: Initial OCR Extraction.</b> We applied optical character recognition to each page image to produce raw text output. This initial extraction captured the broad structure of each page\u2014play text, commentary, footnotes, marginalia\u2014but introduced systematic errors particularly concentrated in the footnote regions where character resolution was lowest."
    ))

    paras.append(body_para(
        "<b>Stage 2: LLM-Augmented Validation and Error Correction.</b> We employed Google Gemini as a contextual validation layer, selected specifically for its large context window capacity, which permitted processing entire pages while maintaining awareness of the relationships between play text, commentary, and footnote apparatus. The model received both the raw OCR output and, where available, the original page image, enabling it to identify and correct character-level errors that standard OCR post-processing heuristics would miss. Gemini proved particularly effective at three categories of correction: (1) resolving abbreviated citations where single-character OCR errors rendered references unidentifiable (e.g., distinguishing \u201cSteev.\u201d from \u201cSteeu.\u201d or \u201cClar.\u201d from \u201cClar,\u201d); (2) accurately reading dense footnote typography where character boundaries were ambiguous; and (3) preserving the complex typographic conventions of historical scholarly editions, including italicized lemma markers, parenthetical attributions, and nested quotation structures. We acknowledge that alternative large-context models might achieve comparable correction rates, and the choice of Gemini was driven primarily by its context window capacity at the time of development rather than an exhaustive comparative evaluation."
    ))

    paras.append(body_para(
        "<b>Stage 3: Structural Parsing and Segmentation.</b> Following validation, the corrected text output for each page was parsed to identify structural boundaries: act and scene divisions, individual line numbers, speaker attributions, stage directions, and discrete commentary entries. This parsing employed a combination of regular expression pattern matching and contextual heuristics informed by the consistent formatting conventions across the Variorum series\u2014for example, scene headings invariably appear in a distinct typographic style, and line numbers occupy fixed positions in the margin. The output of this stage was a semi-structured representation of each volume organized by act, scene, and line number, with associated commentary entries tagged to their corresponding line references."
    ))

    # ── C ──
    paras.append(subsection_header("C. Structured Data Architecture: JSON Schema Design"))

    paras.append(body_para(
        "A critical architectural decision concerned the data format for representing the digitized Variorum corpus within the web platform. We elected to use structured JSON (JavaScript Object Notation) rather than TEI-XML (Text Encoding Initiative) for several reasons that bear on both scholarly methodology and practical implementation. The Variorum\u2019s commentary density and cross-referential complexity exceed what can be practically encoded in TEI without producing unwieldy documents that impede real-time web rendering. Our emphasis on line-level computational alignment requires flexible, hierarchically nested data structures that JSON supports natively. Furthermore, integration with contemporary language model APIs\u2014which accept and return JSON payloads\u2014is simplified by maintaining the scholarly data in the same format throughout the pipeline, eliminating serialization overhead."
    ))

    paras.append(body_para(
        "Each of the 20 digitized plays is stored as an independent JSON file following a consistent schema. The top-level object uses section headings as keys (e.g., \u201cDRAMATIS PERSONAE,\u201d \u201cACT 1 SCENE 1,\u201d \u201cACT 3 SCENE 4\u201d). Within each section, individual line numbers serve as string keys mapping to objects containing two fields: (1) <b>play</b> (string)\u2014the modernized Folger Shakespeare Library text for that line, including speaker prefixes and stage directions where applicable; and (2) <b>notes</b> (array of strings)\u2014zero or more Variorum commentary entries associated with that line, each preserving the original attribution, lemma, and scholarly discussion as a contiguous text block."
    ))

    paras.append(body_para(
        "This schema enables O(1) lookup of both play text and historical commentary for any given act, scene, and line combination\u2014a property essential for the platform\u2019s real-time interaction model. The play-level JSON files range from approximately 12,000 to 15,000 lines depending on the density of commentary for each play, stored in the platform\u2019s Public/Data/ directory and served as static assets."
    ))

    # ── D ──
    paras.append(subsection_header("D. Editorial Alignment: Bridging Early Modern and Contemporary Editions"))

    paras.append(body_para(
        "A fundamental methodological challenge lay in aligning the Variorum\u2019s Early Modern English text\u2014which preserves original folio and quarto spellings\u2014with a contemporary edition suitable for modern readers. We selected the Folger Shakespeare Library digital texts as the target edition for alignment, based on five criteria: (1) <i>Scholarly authority</i>\u2014the Folger editions are produced according to rigorous editorial standards while remaining faithful to early authoritative witnesses; (2) <i>Accessibility</i>\u2014they employ modernized spelling and punctuation that render Shakespeare\u2019s language comprehensible without specialized training in Early Modern English orthography; (3) <i>Referential stability</i>\u2014the Folger digital texts maintain consistent, versioned line numbering that enables reliable programmatic alignment; (4) <i>Open access</i>\u2014the Folger\u2019s commitment to open access aligns with the project\u2019s democratizing mission and avoids licensing constraints; and (5) <i>Editorial philosophy</i>\u2014the Folger\u2019s approach of modernizing orthography while preserving semantic and poetic integrity matches the project\u2019s goal of accessibility without scholarly compromise."
    ))

    paras.append(body_para(
        "The alignment process operated at line granularity using Gemini\u2019s natural language processing capabilities. For each annotated passage in the Variorum, the system identified the corresponding line or lines in the Folger edition by: extracting the Variorum\u2019s line reference (typically a marginal line number or range); normalizing the Early Modern English text to remove spelling variants, archaic letter forms, and typographic anomalies; matching the normalized text against the Folger edition\u2019s corresponding scene, accounting for differences in lineation, speech prefix conventions, and stage direction placement; and recording the alignment as a mapping from the Variorum\u2019s internal reference system to the Folger\u2019s line numbering."
    ))

    paras.append(body_para(
        "This process achieved greater than 95% successful matching across the corpus. The remaining approximately 5% of alignment failures arose primarily from three sources: editorial differences in verse lineation (where the Variorum and Folger disagree on line breaks), passages where the Variorum commentary references text that the Folger omits or significantly emends, and stage directions whose placement differs between editions. These edge cases were flagged for manual review and, where resolvable, corrected by hand."
    ))

    # ── E ──
    paras.append(subsection_header("E. Systematic Bibliographic Resolution via Computational Agents"))

    paras.append(body_para(
        "One of the most significant contributions of this project is the systematic resolution of the Variorum\u2019s notoriously opaque citation practices. The NVS editions employ abbreviated references throughout their commentary\u2014\u201cJohns.\u201d for Samuel Johnson, \u201cMal.\u201d for Edmond Malone, \u201cClar.\u201d for Edward Capell, \u201cMurr.\u201d for James Murray\u2014that presume familiarity with a bibliographic apparatus often located dozens or hundreds of pages from the commentary itself. More problematically, many significant figures\u2014lexicographers, actors, theater managers, continental critics\u2014appear in prose introductions and appendices without ever being included in the formal bibliographic tables, rendering their identification nearly impossible even for specialists."
    ))

    paras.append(body_para(
        "We employed Anthropic\u2019s Claude, accessed through Claude Code, to develop computational agents that processed the entire 20-volume corpus through a four-stage resolution pipeline:"
    ))

    paras.append(body_para(
        "<b>Stage 1: Comprehensive Extraction.</b> Agents systematically scanned every page of each volume\u2014not merely the commentary sections, but also introductory essays, appendices, performance histories, and bibliographic tables\u2014to identify all personal names, initials, and abbreviated references. This exhaustive approach was essential because the Variorum\u2019s citation practices are inconsistent across volumes and editors; a figure abbreviated in one volume may appear in full in another, and cross-volume analysis significantly improves resolution rates."
    ))

    paras.append(body_para(
        "<b>Stage 2: Contextual Classification.</b> Each extracted reference was classified by type (editor, textual critic, actor, theater manager, lexicographer, historical figure) and by the context in which it appeared (textual commentary, performance history, linguistic analysis, bibliographic appendix). This classification served both disambiguation purposes\u2014distinguishing, for example, between multiple figures sharing similar abbreviations\u2014and enrichment purposes, enabling the platform to present each figure with appropriate contextual metadata."
    ))

    paras.append(body_para(
        "<b>Stage 3: Cross-Reference Resolution.</b> The agents compared abbreviated citations across the full corpus, leveraging the redundancy inherent in a 20-volume series covering overlapping scholarly traditions. An abbreviation appearing without resolution in one volume could often be resolved by locating its full-form equivalent in another volume\u2019s bibliography or introductory apparatus. The system employed frequency analysis and contextual clustering to disambiguate cases where multiple scholars shared similar abbreviations."
    ))

    paras.append(body_para(
        "<b>Stage 4: Bibliographic Completion.</b> For each resolved figure, the system generated a comprehensive identification record including their full name, scholarly contributions, active period, institutional affiliations where determinable, and their theoretical or editorial approach as evidenced by their commentary across the corpus."
    ))

    paras.append(body_para(
        "This computational approach catalogued over 3,500 individual critics and scholars, resolved approximately 850 abbreviated citations, identified over 400 cross-references between volumes, and recorded more than 300 performance references with identified actors and productions. Approximately 15% of abbreviated references resisted automated resolution due to insufficient contextual information; these were flagged for manual research, with accessible entries completed through traditional bibliographic methods."
    ))

    # ── F ──
    paras.append(subsection_header("F. Critics and Supplementary Content Extraction Pipeline"))

    paras.append(body_para(
        "Beyond the core Variorum commentary, each play required supplementary content detailing the historical critics whose annotations appear in that volume\u2019s apparatus. This content was sourced from a set of 20 Microsoft Word documents (.docx format), one per play, containing structured biographical and bibliographic information about the New Variorum critics."
    ))

    paras.append(body_para(
        "We developed a Python extraction pipeline that automated the conversion of these documents into web-ready HTML. The pipeline operated through the following stages: (1) <i>DOCX XML Parsing</i>\u2014each .docx file was opened as a ZIP archive, and the underlying word/document.xml was parsed using Python\u2019s zipfile and xml.etree.ElementTree modules, providing direct access to the document\u2019s XML structure without requiring Microsoft Office or third-party DOCX libraries; (2) <i>Style and Formatting Extraction</i>\u2014the parser traversed the XML tree to identify paragraph styles (headings at levels 3, 4, and 5; body text; table structures) and inline formatting runs (bold, italic, combined bold-italic), with each paragraph\u2019s style determined by examining the w:pStyle element and run-level formatting extracted from w:rPr properties; (3) <i>HTML Generation</i>\u2014extracted content was converted to semantically appropriate HTML: headings to h3/h4/h5 tags, body paragraphs to p tags, bold and italic runs to strong and em elements respectively, and tabular data to table/th/td structures; (4) <i>Modal Integration</i>\u2014the generated HTML was injected into play-specific modal dialogs within the platform\u2019s index.html, each with a unique identifier; and (5) <i>Dynamic Routing</i>\u2014the JavaScript function openCriticsModal() was updated with conditional branching logic keyed on the currentPlay global variable to route each play to its corresponding modal."
    ))

    paras.append(body_para(
        "This pipeline was applied to all 20 plays, producing per-play critics modals with content precisely matching the source documents while inheriting the platform\u2019s typographic and layout conventions."
    ))

    # ── G ──
    paras.append(subsection_header("G. Web Platform Architecture and Single-Page Application Design"))

    paras.append(body_para(
        "The Shakespeare Digital Variorum is implemented as a monolithic single-page application (SPA) deployed as a static site with serverless computational functions. This architectural choice was motivated by long-term sustainability considerations: static sites impose minimal operational overhead, reducing ongoing costs and technical maintenance burden; the separation of historical data (JSON files) from computational processing (serverless functions) makes the platform\u2019s methods transparent and auditable; and the lightweight infrastructure supports the project\u2019s commitment to open access."
    ))

    paras.append(body_para(
        "<b>Presentation Layer.</b> The entire user interface is contained within a single index.html file comprising approximately 23,000 lines of HTML, CSS, and JavaScript. The layout employs a three-column grid structure: (1) a left navigation sidebar containing the play selector dropdown and a hierarchical act/scene navigation list enabling users to browse the corpus and navigate within individual plays; (2) a center play text reader displaying the modernized Folger text organized by act and scene, with line-level selection capability where users select lines by clicking, constrained to a maximum of five contiguous lines to bound the scope of computational analysis requests; and (3) a right analysis and commentary panel hosting the tiered analysis interface, historical Variorum commentary display, follow-up question input, and external research links."
    ))

    paras.append(body_para(
        "<b>State Management.</b> Application state is maintained through global JavaScript variables including currentPlay (the active play identifier), currentScene (the active scene), selectedLines (an array of currently selected line objects), and currentTab (the active analysis tier). Play selection triggers loadSelectedPlay(), which fetches the corresponding JSON data file, parses the act/scene structure, populates the navigation list, and renders the first scene\u2019s text in the reader panel."
    ))

    paras.append(body_para(
        "<b>CSS Design Token System.</b> Visual styling is organized through a dedicated CSS custom properties file (styles/tokens.css) that defines semantic design tokens for colors, typography, spacing, and elevation. The token file imports the Inter and Source Serif 4 font families via Google Fonts and defines root-level variables including --bg, --panel, --surface, --ink-1 through --ink-3, --line, --shadow, --hl, and --accent. A prefers-color-scheme: dark media query provides automatic dark mode adaptation by overriding these tokens with appropriate dark palette values. Component-level styles within index.html reference these tokens, ensuring consistent theming throughout the interface."
    ))

    paras.append(body_para(
        "<b>Modal System.</b> Supplementary content is presented through a system of overlay modals, each with dedicated open/close functions and click-outside-to-dismiss behavior. The modal inventory includes: \u201cHow to Use,\u201d \u201cHow to Help,\u201d \u201cAbout,\u201d a \u201cNew Variorum\u201d explainer, \u201cRights,\u201d and 20 play-specific critics modals. Each modal\u2019s visibility is controlled through element.style.display toggling between \u2018flex\u2019 and \u2018none.\u2019"
    ))

    # ── H ──
    paras.append(subsection_header("H. Serverless AI Integration and Tiered Analysis Framework"))

    paras.append(body_para(
        "The platform\u2019s computational analysis capabilities are implemented through Netlify serverless functions that route user requests to large language models based on the selected analysis tier. This architecture separates the static presentation layer from the computational backend, enabling independent scaling and cost management."
    ))

    paras.append(body_para(
        "<b>Primary API Endpoint.</b> The main serverless function (functions/shakespeare.js) exposes a POST endpoint at /api/shakespeare. It accepts a JSON body containing the selected text, analysis level (or mode), playName, and sceneName. The function initializes an OpenAI client and routes requests based on the analysis mode: (1) <i>Basic mode</i> employs GPT-4o-mini with temperature 0.7 and a maximum token budget of 800, generating plain-language paraphrases and essential glosses suitable for readers encountering Shakespeare for the first time; (2) <i>Expert mode</i> employs GPT-4o-mini with temperature 0.7 and a maximum token budget of 3,000, producing analysis that incorporates historical context, critical traditions, and comparative editorial perspectives; and (3) <i>Full Fathom Five mode</i> employs GPT-4o (the full model, upgraded from mini) with temperature 0.3 and a maximum token budget of 8,000, generating comprehensive commentary intended to match the depth and rigor of traditional Variorum annotation. The lower temperature parameter enforces greater determinism and factual precision appropriate for scholarly output."
    ))

    paras.append(body_para(
        "<b>Edge Function (Alternative Path).</b> A Deno-based edge function (netlify/edge-functions/variorum-edge.js) provides an alternative streaming analysis path using Anthropic\u2019s Claude via the Anthropic Messages API. This function returns Server-Sent Events (SSE) with typed chunks (start, section, content, complete, error), enabling progressive rendering of analysis results. The edge function accesses the Claude API via an encrypted environment variable and implements its own note-loading and text-matching logic."
    ))

    paras.append(body_para(
        "<b>Prompt Engineering.</b> Each analysis tier employs a distinct system prompt engineered to produce output appropriate to the target audience. The Full Fathom Five prompt, in particular, instructs the model to assume the role of a Shakespeare scholar producing variorum-style annotation, incorporating textual variants, editorial history, critical traditions, and linguistic analysis. All prompts explicitly identify the play, act, scene, and selected lines to anchor the model\u2019s output in the correct textual context."
    ))

    # ── I ──
    paras.append(subsection_header("I. Retrieval-Augmented Generation for Historical Commentary Delivery"))

    paras.append(body_para(
        "A distinguishing feature of the platform\u2019s architecture is its integration of historical Variorum commentary with generative AI analysis through a retrieval-augmented generation (RAG) approach. Rather than relying solely on the language model\u2019s parametric knowledge of Shakespeare scholarship, the system retrieves and delivers the actual historical notes from the digitized corpus."
    ))

    paras.append(body_para(
        "<b>Server-Side Note Retrieval.</b> When a Full Fathom Five analysis request is received, the serverless function loads the corresponding play\u2019s JSON data file either from the deployed site\u2019s public URL or from GitHub raw content URLs as a fallback. The function findRelevantNotes iterates across all scenes in the JSON structure, normalizing the selected text (removing punctuation variations, standardizing whitespace, and handling stage direction markers) and matching it against the play field of each line entry. When a match is identified, the associated notes array is extracted."
    ))

    paras.append(body_para(
        "<b>Deterministic Retrieval Strategy.</b> The note retrieval mechanism deliberately employs deterministic text matching rather than semantic similarity search. This design decision prioritizes precision and provenance: users can be confident that the historical commentary they receive corresponds exactly to the lines they selected, with no risk of the false positives that embedding-based semantic search might introduce. The trade-off is reduced recall for paraphrased or obliquely related commentary, which we consider acceptable given the platform\u2019s emphasis on scholarly reliability."
    ))

    paras.append(body_para(
        "<b>Separation of Authority.</b> The retrieved historical notes are presented in a dedicated \u201cNew Variorum Analysis\u201d section of the response, clearly separated from the AI-generated analysis sections. This architectural separation maintains a bright line between verified historical scholarship (the Variorum commentary, representing the work of identified critics over centuries) and computational output (the language model\u2019s generated analysis, which is explicitly labeled as machine-generated and subject to the standard limitations of large language models). Both the client and server maintain whitelists (playsWithNewVariorum and playsWithoutNewVariorum respectively) that control the availability of Variorum-specific features based on whether a given play\u2019s JSON data has been fully processed."
    ))

    # ── J ──
    paras.append(subsection_header("J. Geneva Bible Contextual Integration"))

    paras.append(body_para(
        "Recognizing the significance of biblical allusion in Shakespeare\u2019s works and the interpretive tradition surrounding it, the platform integrates the Geneva Bible (1560)\u2014the translation most widely available during Shakespeare\u2019s lifetime and demonstrably influential on his language\u2014as a contextual reference source."
    ))

    paras.append(body_para(
        "The integration is implemented through a dedicated JavaScript class (BibleSearch, defined in Public/Data/bibleSearch.js) that loads and indexes the full Geneva Bible text from a plain-text data file. Upon initialization, the class parses the text into individual verses keyed by book, chapter, and verse number, constructing an in-memory index that supports rapid lookup. The getBibleContext method accepts a text selection and identifies potentially relevant biblical passages through textual correspondence analysis, while the formatBibleContextForPrompt method structures identified biblical references into a format suitable for inclusion in language model prompts, enabling the AI analysis to incorporate biblical context when relevant to the selected passage."
    ))

    # ── K ──
    paras.append(subsection_header("K. External Resource Discovery and Multimedia Integration"))

    paras.append(body_para(
        "The platform facilitates discovery of contemporary scholarship and performance materials through contextually constructed search integrations that leverage the user\u2019s current selection and navigation state:"
    ))

    paras.append(body_para(
        "<b>Performance Discovery.</b> The searchYouTube function constructs search queries combining the current play name, act/scene identifier, and the text \u201cperformance,\u201d opening YouTube search results in a new browser tab. This enables users to discover recordings of how different actors and directors have interpreted the specific lines under examination."
    ))

    paras.append(body_para(
        "<b>Scholarly Research Integration.</b> The searchJSTOR function generates structured academic queries using Boolean AND-clauses: a quoted play title, act and scene identifiers formatted in multiple conventions (Arabic, Roman numeral), and a selected phrase from the highlighted text. Similarly, searchGoogleScholar constructs scoped academic searches with the addition of the term \u201cShakespeare\u201d to improve precision. The openInternetShakespeare function maps the current play to the Internet Shakespeare Editions\u2019 URL slug system and opens the corresponding ISE page, providing direct access to an independent scholarly edition for cross-reference."
    ))

    paras.append(body_para(
        "<b>Audio Integration.</b> For selected plays where audio recordings are available (currently <i>All\u2019s Well That Ends Well</i>), the platform displays an audio player with playback controls including play/pause, seek, volume, and speed adjustment. The player\u2019s visibility is controlled by a polling interval that checks whether the currently selected play matches the available audio inventory and shows or hides the player accordingly."
    ))

    # ── L ──
    paras.append(subsection_header("L. Interactive Question-and-Answer System"))

    paras.append(body_para(
        "Beyond the tiered analysis system, the platform provides an interactive follow-up question capability that enables users to engage in contextual dialogue about selected passages. The Q&A interface presents a text input and submit button; upon submission, the askFollowUp() function constructs a composite context string incorporating the current play name, the originally selected text, the most recent analysis output, and the user\u2019s question."
    ))

    paras.append(body_para(
        "This composite context is transmitted to the same serverless endpoint with the analysis mode set to \u201cfollowup.\u201d The server-side handler employs a dedicated system prompt optimized for conversational scholarly assistance, instructing the model to answer with reference to the specific textual and analytical context rather than generating generic Shakespeare commentary. Responses are rendered as distinct follow-up answer blocks inserted into the analysis panel, preserving the conversational thread while maintaining visual separation from the primary analysis output."
    ))

    # ── M ──
    paras.append(subsection_header("M. Data Quality Assurance and Iterative Correction"))

    paras.append(body_para(
        "The digitization and alignment pipeline required ongoing quality assurance to identify and correct data integrity issues. Several categories of correction were performed iteratively throughout the development process:"
    ))

    paras.append(body_para(
        "<b>Missing Play Text.</b> In some cases, the JSON data files contained complete commentary (the notes array) but empty or missing play text (the play field) for specific scenes. For example, the Hamlet data file was found to be missing play text for Act 1 Scenes 3, 4, and 5, despite the commentary entries being correctly populated. These gaps were remedied by manually transcribing the corresponding Folger text and populating the play fields, preserving the existing notes arrays intact."
    ))

    paras.append(body_para(
        "<b>Content Normalization.</b> The critics and bibliography content extracted from the 20 per-play DOCX files required normalization to ensure consistent formatting across plays. A dedicated Python script standardized heading levels, paragraph spacing, and citation formatting to match the template established by the Macbeth critics modal\u2014the first play processed and subsequently used as the reference standard."
    ))

    paras.append(body_para(
        "<b>JavaScript Logic Debugging.</b> The dynamic routing logic for play-specific modals required careful debugging to ensure correct behavior across all 20 plays. Issues identified and corrected included: duplicate function definitions in the monolithic index.html (where later definitions silently overrode earlier ones), missing entries in the playsWithCriticsButtons array that rendered the critics button invisible for certain plays, and inconsistent modal close behavior requiring the addition of try/catch blocks to handle modals that might not exist in the DOM for certain play states."
    ))

    # ── N ──
    paras.append(subsection_header("N. Deployment Infrastructure and Continuous Delivery"))

    paras.append(body_para(
        "The platform is deployed on Netlify\u2019s static hosting infrastructure, configured through a netlify.toml manifest that specifies the build environment (Node.js version 22), the functions directory, and URL redirect rules. The SPA routing pattern is implemented through a catch-all redirect (/* to /index.html) that enables client-side navigation without server-side route resolution. JSON data files are served with no-cache headers to ensure users always receive the latest digitized content."
    ))

    paras.append(body_para(
        "Serverless functions execute in Netlify\u2019s Node.js runtime environment, with API keys for OpenAI and Anthropic managed through Netlify\u2019s encrypted environment variable system. The edge function operates in Netlify\u2019s Deno-based edge runtime, enabling lower-latency streaming responses from geographically distributed points of presence."
    ))

    paras.append(body_para(
        "Version control is maintained through Git, with the complete codebase\u2014including JSON data files, serverless functions, and the monolithic frontend\u2014hosted in a single repository. This monorepo approach simplifies deployment (Netlify automatically rebuilds and deploys on push to the main branch) at the cost of repository size, which grows with each new play\u2019s JSON data file."
    ))

    return paras


# ── main ─────────────────────────────────────────────────────────────────

def main():
    doc = Document(DOCX_PATH)
    body = doc.element.body

    # Paragraphs 33 through 82 are old Sections III and IV
    OLD_START = 33  # "III. Methodology"
    OLD_END   = 82  # last paragraph of "IV. Technical Implementation" (inclusive)

    # Collect the XML elements we need to remove
    old_elements = [doc.paragraphs[i]._element for i in range(OLD_START, OLD_END + 1)]

    # The element that will come AFTER our new content (paragraph 83 = "V. Results...")
    anchor = doc.paragraphs[OLD_END + 1]._element

    # Remove old paragraphs
    for elem in old_elements:
        body.remove(elem)

    # Build new methodology paragraphs
    new_paras = build_new_methodology()

    # Insert new paragraphs before the anchor (V. Results...)
    anchor_index = list(body).index(anchor)
    for i, p_elem in enumerate(new_paras):
        body.insert(anchor_index + i, p_elem)

    # Renumber Section V -> IV, VI -> V, VII -> VI
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        # Only touch section headers (bold, 18pt)
        if p.runs and p.runs[0].font.size == SECTION_FONT and p.runs[0].bold:
            if text.startswith("V. "):
                p.runs[0].text = text.replace("V. ", "IV. ", 1)
            elif text.startswith("VI. "):
                p.runs[0].text = text.replace("VI. ", "V. ", 1)
            elif text.startswith("VII. "):
                p.runs[0].text = text.replace("VII. ", "VI. ", 1)

    doc.save(DOCX_PATH)
    print("Done. Sections III & IV replaced with expanded methodology.")
    print(f"Old sections: paragraphs {OLD_START}-{OLD_END} (50 paragraphs removed)")
    print(f"New methodology: {len(new_paras)} paragraphs inserted")
    print(f"Sections V/VI/VII renumbered to IV/V/VI")


if __name__ == "__main__":
    main()
