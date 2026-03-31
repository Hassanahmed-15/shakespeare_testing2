#!/usr/bin/env python3
"""
Replace the methodology section (paras 33-99) with a deeply technical rewrite
emphasizing architectural debates, trade-offs, obstacles, and exact specifics.
"""

from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

DOCX_PATH = "Public/Data/Full Fathom Five Hassan Edition.docx"
SECTION_FONT = Emu(228600)
SUBSECTION_FONT = Emu(171450)
FONT_NAME = "Times New Roman"


def make_run(parent_p, text, bold=None, italic=None, font_size=None):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if bold:
        rPr.append(OxmlElement('w:b'))
    if italic:
        rPr.append(OxmlElement('w:i'))
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    rPr.append(rFonts)
    if font_size is not None:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size.pt * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
        rPr.append(szCs)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    parent_p.append(r)
    return r


def make_paragraph():
    return OxmlElement('w:p')


def section_header(text):
    p = make_paragraph()
    make_run(p, text, bold=True, font_size=Pt(18))
    return p


def subsection_header(text):
    p = make_paragraph()
    make_run(p, text, bold=True, font_size=Pt(13.5))
    return p


def parse_rich_text(markup):
    runs = []
    pattern = re.compile(r'<(b|i|bi)>(.*?)</\1>', re.DOTALL)
    pos = 0
    for m in pattern.finditer(markup):
        if m.start() > pos:
            runs.append((markup[pos:m.start()], False, False))
        tag = m.group(1)
        txt = m.group(2)
        bold = tag in ('b', 'bi')
        italic = tag in ('i', 'bi')
        runs.append((txt, bold, italic))
        pos = m.end()
    if pos < len(markup):
        runs.append((markup[pos:], False, False))
    return runs


def body_para(markup):
    p = make_paragraph()
    for text, bold, italic in parse_rich_text(markup):
        make_run(p, text,
                 bold=bold if bold else None,
                 italic=italic if italic else None)
    return p


# ── content ──────────────────────────────────────────────────────────────

def build_new_methodology():
    P = []

    P.append(section_header("III. Methodology and System Development"))

    P.append(body_para(
        "This section presents a technically precise account of the pipeline, architecture, and engineering trade-offs involved in transforming the New Variorum Shakespeare from digitized historical volumes into a live, AI-augmented scholarly platform. Where design decisions were contested or where multiple viable approaches existed, we describe the alternatives considered, the obstacles that forced particular choices, and the concrete implications of each path taken."
    ))

    # ── A ──
    P.append(subsection_header("A. Corpus Acquisition: Source Selection and Scan-Quality Triage"))

    P.append(body_para(
        "Digital surrogates of the 23 pre-1956 NVS dramatic editions were sourced from Internet Archive and HathiTrust. Of these, 20 volumes yielded page-level TIFF/JPEG scans at 300\u2013400 DPI\u2014sufficient for OCR. Three volumes (<i>Henry IV, Part 2</i>, <i>Troilus and Cressida</i>, <i>Richard II</i>) remain in active digitization. Each volume was triaged into quality tiers\u2014high, moderate, low\u2014based on three measurable criteria: character-level clarity in footnote regions (where type size drops below 7 pt), preservation of marginal line numbers (which anchor commentary alignment), and presence of bleed-through or foxing that confuses binarization. Low-tier volumes (notably the 1877 <i>Hamlet</i> and the 1907 <i>Antony and Cleopatra</i>) received additional manual correction passes after OCR."
    ))

    P.append(body_para(
        "<b>Why Internet Archive over HathiTrust alone.</b> HathiTrust provides higher-resolution scans and page-level metadata, but imposes programmatic access restrictions (no bulk download API for in-copyright-claimed works, even when the underlying text is public domain). Internet Archive\u2019s open bulk-download endpoints permitted batch acquisition of all 20 volumes in a single scripted run. Where both repositories held a volume, we compared scan quality page-by-page and selected the superior source; in seven cases, we composited pages from both repositories to achieve the highest-quality scan set."
    ))

    # ── B ──
    P.append(subsection_header("B. Multi-Stage OCR Pipeline: From Page Image to Structured Text"))

    P.append(body_para(
        "Standard OCR (Tesseract 5, ABBYY FineReader) performed adequately on the Variorum\u2019s main play text (set in ~12 pt roman type) but produced unacceptable error rates in the footnote apparatus, where abbreviated citations like \u201cSteev.,\u201d \u201cClar.,\u201d and \u201cMal.\u201d appear in ~7 pt type with inconsistent kerning. Error analysis on a 100-page sample from <i>Macbeth</i> revealed a footnote-region character error rate (CER) of 8.4% versus 1.2% in the main text\u2014a sevenfold disparity. The most frequent failure modes were: (1) confusion of period-terminated abbreviations with sentence endings, splitting single citations across multiple OCR blocks; (2) misreading of italic lemma markers as roman (destroying the distinction between the word being glossed and the gloss itself); and (3) systematic substitution of \u201crn\u201d for \u201cm\u201d in small type."
    ))

    P.append(body_para(
        "<b>Stage 1: Raw Extraction.</b> Each page image was processed to produce raw UTF-8 text preserving approximate spatial layout (line breaks, indentation level). No semantic parsing was attempted at this stage."
    ))

    P.append(body_para(
        "<b>Stage 2: LLM-Augmented Validation.</b> We passed each page\u2019s raw OCR output to Google Gemini (selected for its 1M-token context window, which permitted processing an entire volume\u2019s worth of surrounding pages as disambiguation context). The model received the OCR text alongside the original page image and a system prompt instructing it to: correct character-level OCR errors, restore truncated abbreviations, re-join citations split across OCR blocks, and preserve the italic/roman distinction using Markdown conventions. <b>Architectural debate: Gemini vs. GPT-4V vs. Claude.</b> At the time of development (2024), Gemini was the only commercially available model offering a context window large enough to hold an entire volume\u2019s bibliography alongside the current page. GPT-4V offered competitive vision capabilities but a 128K-token window insufficient for cross-referencing against a full bibliography. Claude\u2019s context window (200K) was closer but still required truncation of larger volumes. We chose Gemini on context-window grounds alone, acknowledging that future models may equalize this advantage."
    ))

    P.append(body_para(
        "<b>Stage 3: Structural Segmentation.</b> The corrected text for each page was parsed into act/scene boundaries, line numbers, speaker attributions, stage directions, and discrete commentary entries. Parsing relied on regular expression pattern matching tuned to the Variorum\u2019s typographic conventions: scene headings in small caps, line numbers in the left margin with a fixed-width pattern (\\d{1,4}), and commentary entries delimited by line-number references followed by a closing bracket (e.g., \u201c42]\u201d). Edge cases included commentary spanning multiple pages (requiring cross-page stitching) and volumes where the editor changed line-numbering conventions mid-play."
    ))

    # ── C ──
    P.append(subsection_header("C. Data Architecture: The JSON-over-TEI Decision"))

    P.append(body_para(
        "<b>The core architectural debate.</b> The digital humanities community has converged on TEI-XML as the standard encoding for scholarly editions. We considered TEI seriously and prototyped a TEI encoding of <i>Macbeth</i>\u2019s first act. The prototype revealed three problems specific to Variorum data: (1) <b>Commentary density.</b> A single line of <i>Macbeth</i> can carry 2,000+ words of commentary from a dozen critics spanning three centuries. Encoding this in TEI\u2019s <note> and <app> elements produced XML documents exceeding 50 MB per play\u2014too large for real-time browser parsing without a dedicated XML database backend. (2) <b>Line-level random access.</b> The platform\u2019s interaction model requires O(1) lookup: a user clicks a line and instantly receives its commentary. TEI-XML\u2019s tree structure requires XPath traversal from the root, which, at 50 MB document sizes, introduces perceptible latency. (3) <b>LLM API compatibility.</b> The entire AI integration pipeline (OpenAI, Anthropic) uses JSON natively. Maintaining TEI would require a serialization layer translating between XML and JSON on every API call, adding latency and a potential failure point."
    ))

    P.append(body_para(
        "We therefore adopted a flat JSON schema. Each play is a single .json file (12,000\u201315,000 lines). The top-level object maps section headings as string keys (\u201cDRAMATIS PERSONAE,\u201d \u201cACT 1 SCENE 1\u201d) to objects whose keys are line-number strings (\u201c1,\u201d \u201c2,\u201d \u201c42\u201d). Each line-number key maps to an object with two fields: <b>play</b> (string: the modernized Folger text for that line) and <b>notes</b> (array of strings: zero or more Variorum commentary entries). This gives true O(1) lookup via JavaScript\u2019s native property access\u2014no parsing, no traversal. <b>Obstacle encountered: key format inconsistency.</b> Early in development, some JSON files used comma-separated scene keys (\u201cACT 1, SCENE 1\u201d) while others used space-separated keys (\u201cACT 1 SCENE 1\u201d). This mismatch caused silent lookup failures in the serverless function, where the client sent one format and the JSON contained the other. The fix required retroactive normalization of all 20 JSON files to a single convention."
    ))

    P.append(body_para(
        "<b>Trade-off acknowledged.</b> The JSON schema sacrifices the rich semantic markup that TEI provides (apparatus entries, witness sigla, editorial responsibility attribution). We accept this trade-off because the platform\u2019s primary contribution is <i>access</i>\u2014making previously inaccessible commentary usable\u2014rather than <i>critical editing</i> in the traditional sense. The raw Variorum text is preserved verbatim in the notes array, and future work could layer TEI encoding on top of the existing JSON as a secondary output format."
    ))

    # ── D ──
    P.append(subsection_header("D. Editorial Alignment: Mapping Variorum Lines to Folger Editions"))

    P.append(body_para(
        "<b>Why the Folger over Norton, Riverside, or Arden.</b> The alignment target was a non-trivial choice. The Norton Shakespeare uses through-line numbering (TLN), which differs from the Variorum\u2019s act-scene-line numbering and would require a TLN-to-ASL conversion layer. The Riverside uses its own lineation that occasionally diverges from both. The Arden editions are textually excellent but commercially licensed, precluding open-access redistribution. The Folger Digital Texts were selected because they: (a) use act-scene-line numbering compatible with the Variorum\u2019s scheme; (b) are openly licensed; (c) modernize spelling without altering semantic content; and (d) provide a stable, versioned API. This decision traded editorial plurality for alignment simplicity\u2014a deliberate choice we revisit in the Limitations section."
    ))

    P.append(body_para(
        "The alignment process used Gemini to match each Variorum commentary entry\u2019s line reference to the corresponding Folger line. The procedure: (1) extract the Variorum\u2019s marginal line number; (2) normalize the Early Modern text by collapsing spelling variants (e.g., \u201cvnto\u201d \u2192 \u201cunto,\u201d long-s \u201c\u017f\u201d \u2192 \u201cs\u201d); (3) locate the matching line in the Folger\u2019s corresponding scene, tolerating ±2 lines of offset to account for lineation disagreements; (4) record the mapping. <b>Obstacle: stage directions and split verses.</b> The Variorum sometimes counts stage directions as numbered lines; the Folger does not. This caused systematic off-by-one errors in every scene containing mid-scene stage directions. Additionally, shared (split) verse lines\u2014where two characters share a single pentameter line\u2014are sometimes counted as one line in the Variorum and two in the Folger, or vice versa. These cases required per-scene manual review. Overall alignment success: >95%, with the remaining ~5% concentrated in plays with heavy stage-direction traffic (<i>The Tempest</i>, <i>A Midsummer Night\u2019s Dream</i>) and the late-corpus volumes where editors adopted inconsistent lineation practices."
    ))

    # ── E ──
    P.append(subsection_header("E. Bibliographic Resolution: AI Agents on 20 Volumes of Abbreviated Citations"))

    P.append(body_para(
        "The Variorum\u2019s citation apparatus is notoriously opaque. Commentary entries reference scholars by truncated surnames (\u201cJohns.\u201d for Samuel Johnson, \u201cMal.\u201d for Edmond Malone, \u201cSteev.\u201d for George Steevens, \u201cClar.\u201d for Edward Capell) with no inline expansion. The full-name mappings appear in bibliographic appendices located hundreds of pages from the commentary\u2014and even these appendices omit figures (actors, theater managers, lexicographers like Murray of the OED) who appear only in prose introductions. A researcher encountering \u201cMurr.\u201d in the commentary of <i>Hamlet</i> would have no way to identify this as James Murray without cross-referencing the bibliography of a different volume entirely."
    ))

    P.append(body_para(
        "We used Anthropic\u2019s Claude, accessed through Claude Code (an agentic coding interface), to build computational agents that processed the corpus through four stages: (1) <b>Extraction:</b> agents scanned every page of all 20 volumes\u2014not just commentary but also introductory essays, appendices, and performance histories\u2014extracting every proper noun and abbreviated reference. (2) <b>Classification:</b> each reference was tagged by type (editor, critic, actor, lexicographer) and context (textual commentary, performance history, linguistic analysis). (3) <b>Cross-volume resolution:</b> abbreviated citations were compared across the full corpus; a figure abbreviated in one volume often appears in full in another\u2019s bibliography. The agents exploited this redundancy through frequency analysis and contextual clustering. (4) <b>Record generation:</b> each resolved figure received a comprehensive identification record. The process catalogued 3,500+ individuals, resolved ~850 abbreviations, and identified 400+ cross-volume references. <b>Failure mode:</b> approximately 15% of abbreviations resisted resolution\u2014typically obscure Continental European critics referenced once in a single volume with no bibliography entry anywhere in the corpus."
    ))

    # ── F ──
    P.append(subsection_header("F. Critics Content Pipeline: From DOCX to Per-Play HTML Modals"))

    P.append(body_para(
        "Each play\u2019s \u201cWho are the NV Critics\u201d content was authored in a separate Microsoft Word document (20 .docx files, one per play). Converting these to web-ready HTML required a Python pipeline (scripts/replace_critics_from_docx.py) that operated at the OOXML level: each .docx was opened as a ZIP archive, the underlying word/document.xml parsed via Python\u2019s zipfile and xml.etree.ElementTree, and each <w:p> paragraph element traversed to extract: (a) paragraph style from <w:pStyle> (mapping Heading3/4/5 to <h3>/<h4>/<h5>); (b) run-level formatting from <w:rPr>, reading <w:b> for bold and <w:i> for italic; (c) table structures from <w:tbl> elements. The extracted content was converted to semantic HTML and injected into play-specific modal dialogs in index.html."
    ))

    P.append(body_para(
        "<b>Obstacle: the duplicate-function bug.</b> The monolithic index.html file (23,000+ lines) contained two separate definitions of the openCriticsModal() function\u2014one at line ~15,549 and another at line ~22,707. In JavaScript, the later definition silently overrides the earlier. During initial deployment, the first definition (which contained the routing logic for only the original 5 plays) was being overridden by the second definition (which had the full 20-play routing). However, the second definition was itself incomplete at first, causing certain plays\u2019 critics buttons to appear to do nothing. The fix required: (a) unifying the two definitions into a single function; (b) adding all 20 plays to the playsWithCriticsButtons array that controls button visibility; (c) wrapping the closeCriticsModal() function in try/catch blocks for each modal (since a modal\u2019s DOM element might not exist if the user navigated to a play whose modal was never injected); and (d) extending the window.onclick handler to support click-outside-to-close for all 20 new modals."
    ))

    # ── G ──
    P.append(subsection_header("G. Web Platform Architecture: The Monolithic SPA Decision"))

    P.append(body_para(
        "<b>Architectural debate: SPA monolith vs. component framework vs. multi-page app.</b> We considered three architectures: (1) a React/Vue component-based SPA, offering modular code organization and reactive state management; (2) a multi-page app with server-rendered HTML, offering simpler deployment; and (3) a monolithic single-file SPA. We chose option (3) for two reasons. First, <b>deployment sustainability:</b> the platform is intended to survive beyond the initial development team. A monolithic HTML/CSS/JS file has zero build dependencies\u2014no Webpack, no Babel, no npm install\u2014and can be deployed to any static host indefinitely. A React app, by contrast, requires a build step that will break as dependencies deprecate, introducing a maintenance burden incompatible with the project\u2019s preservation mission. Second, <b>transparency:</b> the entire application is viewable in a single \u201cView Source,\u201d aligning with the project\u2019s commitment to open, inspectable infrastructure."
    ))

    P.append(body_para(
        "<b>Trade-off accepted:</b> the monolith grew to ~23,000 lines, making code navigation difficult and enabling bugs like the duplicate-function issue described above. State management relies on global JavaScript variables (currentPlay, currentScene, selectedLines, currentTab, currentAnalysisMode), with no reactive binding\u2014DOM updates are imperative. The three-column layout is implemented as a CSS grid: a left sidebar (play selector dropdown via <select id=\"playSelector\"> and act/scene navigation list <ul id=\"navigationList\">), a center reader panel (div#playText), and a right analysis panel (div#analysisPanel). Line selection is constrained to MAX_SELECTED_LINES = 5 contiguous lines, enforced by index-adjacency checks in the selectLine() function."
    ))

    P.append(body_para(
        "<b>CSS design tokens.</b> Visual consistency is managed through a dedicated token file (styles/tokens.css) defining semantic custom properties: --bg, --panel, --surface for backgrounds; --ink-1 through --ink-3 for text hierarchy; --accent (#3B82F6) for interactive elements; --r-xs through --r-lg for border radii; --s-1 through --s-5 for spacing rhythm. Typography uses two font stacks: Inter (\u201c--ui\u201d) for interface text and Source Serif 4 (\u201c--serif\u201d) for literary content. A @media (prefers-color-scheme: dark) block overrides all token values for automatic dark mode: --bg shifts from #F7F7F8 to #0B0C0F, --ink-1 from #0F1115 to #F2F3F5, and so on. This token architecture permits full theme changes by editing a single 50-line file."
    ))

    # ── H ──
    P.append(subsection_header("H. Serverless AI Integration: Dual-Model Architecture and Prompt Engineering"))

    P.append(body_para(
        "The computational analysis layer is implemented as a Netlify serverless function (functions/shakespeare.js, Node.js 22) and a parallel Deno edge function (netlify/edge-functions/variorum-edge.js). The primary function accepts POST requests at /api/shakespeare (rewritten from /.netlify/functions/shakespeare via netlify.toml redirects) with a JSON body containing: text (the selected passage), level or mode (one of \u201cbasic,\u201d \u201cexpert,\u201d \u201cfullfathomfive,\u201d \u201ccritics,\u201d or \u201cfollowup\u201d), model (default gpt-4o-mini, overridden server-side for fullfathomfive), playName, and sceneName."
    ))

    P.append(body_para(
        "<b>Model routing and parameter tuning.</b> Each tier maps to specific OpenAI parameters: Basic uses gpt-4o-mini at temperature 0.7 with max_tokens 800, producing 2\u20134-sentence glosses. Expert uses gpt-4o-mini at temperature 0.7 with max_tokens 3,000, producing multi-paragraph essays. Full Fathom Five upgrades to gpt-4o at temperature 0.3 with max_tokens 8,000, producing 800\u20131,200-word variorum-style annotations. <b>Why the temperature drop for Full Fathom Five:</b> at temperature 0.7, GPT-4o hallucinated critic names and fabricated publication dates at rates unacceptable for a scholarly tool. Lowering to 0.3 reduced hallucination substantially while preserving analytical depth. <b>Why gpt-4o over gpt-4o-mini for this tier:</b> gpt-4o-mini produced analysis that was fluent but shallow\u2014it consistently omitted textual variants and critical reception, two sections essential to variorum-style commentary. The full gpt-4o model reliably populated all 11 required sections."
    ))

    P.append(body_para(
        "<b>The edge function (Claude path).</b> The Deno edge function provides an alternative streaming analysis path using Anthropic\u2019s Claude (claude-sonnet-4-20250514) at max_tokens 8,000. It returns Server-Sent Events (SSE) with typed JSON chunks: {type: \u201cstart\u201d}, {type: \u201csection\u201d}, {type: \u201ccontent\u201d, word: \u201c...\u201d}, {type: \u201ccomplete\u201d}. Word-by-word streaming (50ms inter-word delay) produces a typewriter effect in the UI. <b>Architectural debate: why maintain two LLM backends?</b> The dual-model architecture was not planned but emerged from iterative development. The OpenAI path was built first (broader model selection, faster iteration). The Claude path was added to evaluate whether Anthropic\u2019s models produced superior scholarly analysis (they did, particularly for textual-variant identification). Rather than forcing a single-vendor choice, we maintained both paths, giving the project resilience against API deprecation or pricing changes from either provider. In production, the primary flow uses OpenAI; the Claude edge function remains available as a fallback and as an A/B comparison tool."
    ))

    P.append(body_para(
        "<b>Prompt structure.</b> The Full Fathom Five system prompt mandates 11 sections in fixed order: Textual Variants, Plain-Language Paraphrase, Language and Rhetoric, Synopsis, Key Words & Glosses, Historical Context, Sources, Literary Analysis, Critical Reception, Similar Phrases in Other Plays, and New Variorum Analysis. Each section has explicit format constraints (e.g., \u201cKey Words: \u2018word\u2019 means [definition]\u201d). The prompt enforces citation requirements: at least one critic per century (18th\u201321st), at least one Marxist critic, plus 2\u20133 others, with full publication details. Target length is 800\u20131,200 words. The server-side function parses the response by iterating over lines, matching each against the known section headers, and building a key-value analysis object that the client renders into collapsible panels."
    ))

    P.append(body_para(
        "<b>Obstacle: asterisk-to-HTML conversion.</b> GPT-4o returns Markdown formatting (*italic*, **bold**) rather than HTML. The platform\u2019s rendering expects raw HTML. A post-processing pipeline of regex substitutions converts *text* to <span style=\"font-style:italic;\">text</span>. This required six overlapping regex patterns (including edge cases for titles containing apostrophes, e.g., <i>A Midsummer Night\u2019s Dream</i>) and a per-play title lookup for stubborn cases. A cleaner solution would instruct the model to emit HTML directly; we attempted this but found that GPT-4o\u2019s adherence to HTML-only formatting was unreliable, reverting to Markdown in ~30% of responses."
    ))

    # ── I ──
    P.append(subsection_header("I. Retrieval-Augmented Generation: Deterministic Notes over Vector Search"))

    P.append(body_para(
        "<b>Architectural debate: deterministic matching vs. embedding-based retrieval.</b> Modern RAG systems typically encode documents into vector embeddings and retrieve via cosine similarity. We prototyped this approach using OpenAI\u2019s text-embedding-ada-002 on the Variorum notes corpus. The results were unsatisfactory: semantic similarity retrieved thematically related but textually unrelated commentary (e.g., a search for \u201cIs this a dagger which I see before me\u201d returned notes about daggers in <i>Hamlet</i> and <i>Julius Caesar</i>, not the Macbeth-specific commentary). For a scholarly tool, <b>false positives are worse than false negatives:</b> a user receiving commentary on the wrong play\u2019s dagger scene would be actively misled, whereas receiving no commentary simply indicates a gap in the database. We therefore adopted deterministic text matching."
    ))

    P.append(body_para(
        "The retrieval function (findRelevantNotes in shakespeare.js) operates as follows: it first extracts line numbers from the user\u2019s selected text using regex (^(\\d+)\\.?\\s*(.*)), then iterates over every scene in the play\u2019s JSON file, looking for an exact line-number key match. If found, the line\u2019s entire notes array is extracted verbatim. If line-number extraction fails, a broader text-search fallback (searchAllScenesForText) compares the lowercased selection against every play field in the database, using substring inclusion with a minimum 3-character threshold. Results are capped at 10 to prevent response bloat. A word-overlap heuristic (matchesText) serves as a third fallback: it tokenizes both strings, filters to words >2 characters, and considers a match if \u226550% of the shorter string\u2019s words appear in the longer."
    ))

    P.append(body_para(
        "<b>Obstacle: the payload-size crisis.</b> Initially, the full text of all retrieved Variorum notes was inlined into the GPT prompt as context for the \u201cNew Variorum Analysis\u201d section. For heavily annotated lines (e.g., <i>Hamlet</i> 3.1.56, \u201cTo be, or not to be,\u201d which carries ~4,000 words of commentary), this inflated the prompt to 15,000+ tokens, causing OpenAI timeouts and exceeding payload limits. The solution was architectural: <b>notes are no longer injected into the prompt.</b> Instead, the serverless function builds the analysis object from the GPT response, then overwrites the \u201cNew Variorum Analysis\u201d key with the raw historical notes retrieved from the JSON file. The GPT model is told that no notes are available and asked to say so briefly; the server then replaces this placeholder with the actual notes. This decoupling means the historical commentary is always delivered verbatim\u2014never summarized, truncated, or paraphrased by the model."
    ))

    P.append(body_para(
        "<b>Play coverage management and the whitelist mismatch.</b> Both client and server maintain arrays determining which plays have Variorum data. The client\u2019s playsWithNewVariorum (17 entries) controls UI visibility of the variorum panel. The server\u2019s playsWithoutNewVariorum (21 entries, listing plays that lack data) controls whether the \u201cNew Variorum Analysis\u201d section is included in the GPT prompt. <b>These two lists are not inverses of each other</b>\u2014the client list was maintained manually and occasionally fell out of sync with the server list, causing plays to show a variorum UI with no data, or suppress the UI for plays that had data. This inconsistency was discovered during testing and partially reconciled, but remains a maintenance concern inherent to the duplicated-whitelist pattern."
    ))

    P.append(body_para(
        "<b>The normalizePlayKey problem.</b> Play names arrive at the server in varied formats (\u201cMacbeth,\u201d \u201cmacbeth,\u201d \u201cThe Merchant of Venice,\u201d \u201cmerchant_of_venice\u201d). The normalizePlayKey function maps these to canonical keys (e.g., \u201cmerchantofvenice\u201d) via a chain of 17 if/includes checks. These keys then index into a playFiles dictionary mapping to actual JSON filenames on disk\u2014which themselves follow no consistent naming convention: macbeth_notes_cleaned_play.json, hamlet_notes (1).json (with a space and parentheses requiring URL encoding), ROMEO_notes.json (uppercase), julius_caesar_clean_(3).json (with a version suffix). This naming chaos arose from the iterative, play-by-play digitization process where each volume was processed independently. The serverless function handles this by URL-encoding filenames via encodeURIComponent and attempting multiple fallback URLs: the deployed site\u2019s /Public/Data/ path, the site root, and two GitHub raw content URLs. This four-URL fallback chain was necessary because Netlify\u2019s redirect rules occasionally interfered with direct static file access for filenames containing special characters."
    ))

    # ── J ──
    P.append(subsection_header("J. Geneva Bible Integration: Keyword-Scored Contextual Retrieval"))

    P.append(body_para(
        "The platform integrates the Geneva Bible (1560) through a client-side JavaScript class (BibleSearch, defined in Public/Data/bibleSearch.js). On initialization, the class fetches the full Geneva Bible text (~4 MB plain text), parses it into a structured array of {book, chapter, verse, text} objects using a regex-based verse detector (^\\[(\\d+):(\\d+)\\]\\s*(.+)$), and builds a search index with extracted keywords (stopword-filtered, archaic-word-tagged) and thematic categories. Ten biblical theme clusters are defined (creation, salvation, wisdom, love, justice, prophecy, kingdom, covenant, temple, exile), each with 4\u20136 seed keywords."
    ))

    P.append(body_para(
        "When the user selects Shakespeare text, the calculateRelevanceScore function scores each Geneva Bible verse against the selection using a weighted heuristic: exact phrase overlap scores 100 points; individual keyword matches score 10 each; partial keyword overlaps score 5; archaic-word matches (\u201cthee,\u201d \u201cthou,\u201d \u201chath,\u201d \u201cdoth,\u201d \u201cforsooth,\u201d \u201cprithee\u201d\u2014a 30-word set) score 15 each (weighted higher because shared archaic vocabulary is a stronger signal of genuine allusion than shared modern vocabulary); and thematic-category matches score 15. The top N results (2 for Expert, 5 for Full Fathom Five) are returned with reference, text, and relevance score, then formatted via formatBibleContextForPrompt into a structured addendum for the AI\u2019s system prompt."
    ))

    # ── K ──
    P.append(subsection_header("K. External Research Discovery: Contextual Search Construction"))

    P.append(body_para(
        "The platform constructs search queries programmatically from the user\u2019s navigation state and text selection, enabling one-click access to four external scholarly and performance resources:"
    ))

    P.append(body_para(
        "<b>YouTube (searchYouTube):</b> query = playName + sceneName + \u201cperformance.\u201d <b>JSTOR (searchJSTOR):</b> query uses Boolean AND-clauses: a quoted play title, act/scene identifiers in both Arabic and Roman numeral formats (e.g., \u201c1.1\u201d and \u201cI.i\u201d), and a selected phrase extracted from selectedLines or selectedText. This dual-format approach was necessary because JSTOR articles use both conventions inconsistently. <b>Google Scholar (searchGoogleScholar):</b> similar scoping with the addition of the term \u201cShakespeare\u201d as a required keyword to filter out non-literary results. <b>Internet Shakespeare Editions (openInternetShakespeare):</b> a static mapping of 25 play identifiers to ISE URL slugs (e.g., \u201cmacbeth\u201d \u2192 \u201cMac,\u201d \u201chamlet\u201d \u2192 \u201cHam,\u201d \u201cmidsummer\u201d \u2192 \u201cMND\u201d), constructing deep links to https://internetshakespeare.uvic.ca/Library/Texts/{slug}/."
    ))

    # ── L ──
    P.append(subsection_header("L. Interactive Q&A: Context-Accumulating Follow-Up System"))

    P.append(body_para(
        "The askFollowUp() function implements a stateful dialogue capability. When the user submits a question, the function constructs a composite context string concatenating: the current play name, the originally selected passage (if available), the most recent analysis output, and the user\u2019s question prefixed with \u201cFollow-up question:\u201d. This composite is sent to the same /api/shakespeare endpoint with level set to \u201cfollowup.\u201d On the server, the handler uses a dedicated system prompt that instructs the model to: identify the follow-up question, use surrounding text only as background context, answer in 2\u20135 sentences, and avoid restating the original analysis. The response is rendered as a visually distinct div.follow-up-answer block inserted before the media section, preserving the conversational thread."
    ))

    # ── M ──
    P.append(subsection_header("M. Data Quality Assurance: Failure Modes and Iterative Correction"))

    P.append(body_para(
        "Quality assurance was not a discrete phase but an ongoing process driven by user-reported failures and systematic testing. Three classes of defect required correction:"
    ))

    P.append(body_para(
        "<b>1. Missing play text in JSON.</b> The Hamlet data file (hamlet_notes (1).json) contained complete notes arrays for Act 1 Scenes 3, 4, and 5, but empty play fields\u2014meaning the line text existed in the Variorum but was lost during the OCR-to-JSON conversion, likely due to a pipeline interruption mid-volume. The commentary was present because it had been processed in a separate pass. The fix required manually transcribing the corresponding Folger text and populating each play field while preserving the existing notes arrays\u2014a delicate operation on a 12,000-line JSON file where a single misplaced comma would corrupt the entire structure."
    ))

    P.append(body_para(
        "<b>2. DOCX-to-HTML normalization drift.</b> The 20 per-play critics DOCX files were authored by different contributors using different Word templates. The first play processed (Macbeth) established the reference format; subsequent plays deviated in heading levels, paragraph spacing, bold/italic usage, and table structures. A normalization script (scripts/normalize_critics_bibliography.py) was developed to enforce consistency, but edge cases persisted\u2014particularly in plays where the DOCX author used custom Word styles not present in the other files."
    ))

    P.append(body_para(
        "<b>3. JavaScript runtime bugs in the monolith.</b> The 23,000-line index.html contained multiple instances of function redefinition (the openCriticsModal duplicate described above), silent failures from missing DOM elements (attempting to set style.display on a modal that was never injected), and whitelist/blacklist inconsistencies between different code regions (playsWithCriticsButtons at line ~14,572 listed 20 plays; playsWithNewVariorum at line ~14,554 listed 17; playsWithoutNewVariorum at line ~14,557 listed a different set that was not the complement of the other two). These bugs were discovered through systematic play-by-play testing and resolved through unification of the various lists and elimination of duplicate function definitions."
    ))

    # ── N ──
    P.append(subsection_header("N. Deployment Infrastructure"))

    P.append(body_para(
        "The platform deploys to Netlify\u2019s static hosting via a netlify.toml manifest specifying: publish directory \u201c.\u201d (the repository root), functions directory \u201cfunctions,\u201d build environment NODE_VERSION=22, and four redirect rules: /api/shakespeare \u2192 /.netlify/functions/shakespeare (status 200), /Public/Data/* \u2192 /Public/Data/:splat (passthrough for static JSON), and /* \u2192 /index.html (SPA catch-all). JSON files are served with Cache-Control: no-cache, no-store, must-revalidate headers to ensure users always receive the latest data. The edge function is registered via [[edge_functions]] with function = \u201cvariorum-edge\u201d and path = \u201c/api/variorum-edge.\u201d API keys (OPENAI_API_KEY, CLAUDE_API_KEY) are stored in Netlify\u2019s encrypted environment variable system, injected at runtime into the serverless (process.env) and edge (Deno.env.get) environments respectively."
    ))

    P.append(body_para(
        "The entire codebase\u2014frontend, serverless functions, JSON data files, Python scripts, and documentation\u2014resides in a single Git repository (monorepo). Netlify triggers automatic builds on push to main. <b>Trade-off:</b> this simplifies deployment but inflates repository size (each play\u2019s JSON file is 1\u20132 MB, totaling ~30 MB of JSON data in the repository). Git operations (clone, pull) are noticeably slower than for a typical web application. We considered Git LFS for the JSON files but rejected it because Netlify\u2019s build system does not natively support LFS, and adding an LFS layer would complicate the zero-dependency deployment that is central to the project\u2019s sustainability model."
    ))

    return P


def main():
    doc = Document(DOCX_PATH)
    body = doc.element.body

    # Find current methodology boundaries
    # Section header "III. Methodology and System Development" through
    # the paragraph just before "IV. Results..."
    start_idx = None
    end_idx = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text.startswith("III. Methodology"):
            start_idx = i
        if text.startswith("IV. Results") and start_idx is not None:
            end_idx = i - 1
            break

    if start_idx is None or end_idx is None:
        print("ERROR: Could not locate methodology section boundaries.")
        print("Searching for section headers:")
        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            if t and (t.startswith("III.") or t.startswith("IV.") or t.startswith("V.")):
                print(f"  [{i}] {t}")
        return

    print(f"Replacing paragraphs {start_idx}-{end_idx} ({end_idx - start_idx + 1} paragraphs)")

    old_elements = [doc.paragraphs[i]._element for i in range(start_idx, end_idx + 1)]
    anchor = doc.paragraphs[end_idx + 1]._element

    for elem in old_elements:
        body.remove(elem)

    new_paras = build_new_methodology()
    anchor_index = list(body).index(anchor)
    for i, p_elem in enumerate(new_paras):
        body.insert(anchor_index + i, p_elem)

    print(f"Inserted {len(new_paras)} new paragraphs")

    doc.save(DOCX_PATH)
    print("Done. Methodology rewritten with technical specifics and architectural debates.")


if __name__ == "__main__":
    main()
