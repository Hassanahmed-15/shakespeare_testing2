#!/usr/bin/env python3
"""
Generate a DOCX file with two columns:
- Column 1: List of all Shakespeare plays
- Column 2: Empty error blocks (one per play)
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# All Shakespeare plays with their full names (from extract_all_plays.py)
PLAYS_CONFIG = {
    "allswell": "All's Well That Ends Well",
    "asyoulikeit": "As You Like It",
    "comedy_errors": "The Comedy of Errors",
    "cymbeline": "Cymbeline",
    "lll": "Love's Labour's Lost",
    "measure": "Measure for Measure",
    "merry_wives": "The Merry Wives of Windsor",
    "merchant": "The Merchant of Venice",
    "midsummer": "A Midsummer Night's Dream",
    "much_ado": "Much Ado About Nothing",
    "pericles": "Pericles, Prince of Tyre",
    "taming_shrew": "The Taming of the Shrew",
    "tempest": "The Tempest",
    "troilus_cressida": "Troilus and Cressida",
    "twelfth_night": "Twelfth Night",
    "two_gentlemen": "Two Gentlemen of Verona",
    "winters_tale": "The Winter's Tale",
    "1henryiv": "Henry IV, Part 1",
    "2henryiv": "Henry IV, Part 2",
    "henryv": "Henry V",
    "1henryvi": "Henry VI, Part 1",
    "2henryvi": "Henry VI, Part 2",
    "3henryvi": "Henry VI, Part 3",
    "henryviii": "Henry VIII",
    "john": "King John",
    "richardii": "Richard II",
    "richardiii": "Richard III",
    "cleopatra": "Antony and Cleopatra",
    "coriolanus": "Coriolanus",
    "hamlet": "Hamlet",
    "julius_caesar": "Julius Caesar",
    "lear": "King Lear",
    "macbeth": "Macbeth",
    "othello": "Othello",
    "romeo_juliet": "Romeo and Juliet",
    "timon": "Timon of Athens",
    "titus": "Titus Andronicus",
}

def generate_plays_errors_docx():
    """Generate DOCX file with plays and error columns"""
    
    # Create a new Document
    document = Document()
    
    # Add title
    title = document.add_heading('Shakespeare Plays - Error Tracking', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a paragraph for spacing
    document.add_paragraph()
    
    # Get all play names and sort alphabetically
    play_names = sorted(PLAYS_CONFIG.values())
    
    # Create table with 38 rows (1 header + 37 plays) and 2 columns
    table = document.add_table(rows=38, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Set column widths (50% each)
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(3.5)
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Play Name'
    header_cells[1].text = 'Errors'
    
    # Format header row - make it bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Populate first column with play names (sorted alphabetically)
    # Leave second column empty for error entries
    for i, play_name in enumerate(play_names, start=1):
        row = table.rows[i]
        row.cells[0].text = play_name
        # Second column (errors) is left empty
        row.cells[1].text = ''
    
    # Add some spacing at the end
    document.add_paragraph()
    
    # Save the document
    output_filename = 'plays_errors_template.docx'
    document.save(output_filename)
    print(f"✅ Successfully created {output_filename}")
    print(f"   - Total plays: {len(play_names)}")
    print(f"   - Table rows: 38 (1 header + 37 plays)")
    print(f"   - Column 1: Play names (sorted alphabetically)")
    print(f"   - Column 2: Empty error blocks (ready for manual entry)")

if __name__ == '__main__':
    generate_plays_errors_docx()
