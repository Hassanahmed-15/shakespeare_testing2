#!/usr/bin/env python3
"""
Generic script to create DOCX files from Shakespeare play JSON.
Handles all scene key formats (ACT 1 SCENE 1, ACT 1,SCENE 1, ACT 1, SCENE 1, etc.)
Includes DRAMATIS PERSONAE and ALL lines with no skipping.

Usage:
  python create_play_docx.py twelfth_night "Twelfth Night"
  python create_play_docx.py henry_iv_part1 "Henry IV, Part 1"
  python create_play_docx.py macbeth_notes_cleaned_play "Macbeth"
"""
import json
import os
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Words that appear in verse/prose but NOT in real speaker names.
# If ANY word in the caps portion before ':' matches, it is NOT a speaker name.
SPEAKER_STOP_WORDS = {
    # Articles / determiners
    'A', 'AN', 'THE', 'SOME', 'SUCH', 'THESE', 'THOSE', 'THIS', 'THAT',
    # Conjunctions
    'AND', 'OR', 'BUT', 'FOR', 'NOR', 'YET', 'SO', 'THAN', 'AS', 'IF', 'LIKE',
    'THEN', 'WHEN', 'WHILE', 'THOUGH',
    # Auxiliary / common verbs
    'IS', 'ARE', 'WAS', 'WERE', 'WILL', 'SHALL', 'CAN', 'MAY', 'MIGHT',
    'DO', 'DOES', 'DID', 'HAS', 'HAVE', 'HAD', 'BE', 'BEEN', 'BEING',
    'MUST', 'SHOULD', 'WOULD', 'COULD',
    'GIVE', 'GAVE', 'GET', 'GOT', 'GO', 'COME', 'CAME', 'TAKE', 'TOOK',
    'KNOW', 'KNEW', 'MAKE', 'MADE', 'SEE', 'SAW', 'SAY', 'SAID', 'TELL',
    'TOLD', 'SHOW', 'FEEL', 'FELT', 'WANT', 'NEED', 'HEAR', 'HEARD',
    'LIVE', 'SEEM', 'LOOK', 'KEEP', 'KEPT', 'CALL', 'HOLD', 'HELD',
    'STAND', 'STOOD', 'FALL', 'FELL', 'RUN', 'RAN',
    # Common dialogue words mistaken as speakers (verse fragments)
    'SHOW', "SHOW'D", 'SHOWED', 'REDOUBLED', 'DOUBLY', 'UPON', 'LIPS',
    'WHORE', 'REBEL', "REBEL'S", 'SKINNY', 'STROKES', 'FOE',
    # Adverbs / negatives
    'NOT', 'NO', 'NEVER', 'EVER', 'MORE', 'MOST', 'LESS', 'LEAST',
    'NOW', 'THEN', 'HENCE', 'THEREFORE', 'THUS',
    # Prepositions
    'IN', 'ON', 'AT', 'BY', 'TO', 'OF', 'WITH', 'FROM', 'INTO', 'UP', 'OUT',
    'UPON', 'UNDER', 'OVER', 'THROUGH', 'BETWEEN', 'AGAINST', 'TOWARD',
    'BEFORE', 'AFTER', 'AMONG', 'WITHIN', 'WITHOUT',
    # Pronouns (modern and archaic)
    'HIS', 'HER', 'THEIR', 'OUR', 'YOUR', 'MY', 'ITS',
    'THEM', 'HIM', 'US', 'WE', 'YOU', 'I', 'HE', 'SHE', 'IT', 'ME',
    'THOU', 'THY', 'THEE', 'YE', 'HIMSELF', 'HERSELF', 'ITSELF',
    'THEMSELVES', 'OURSELVES', 'YOURSELF',
    # Question words
    'WHAT', 'WHO', 'HOW', 'WHICH', 'WHERE', 'WHEN', 'WHY',
}


def is_speaker_name(candidate):
    """
    Return True if candidate (the text before ':') looks like a real speaker name.
    Rules:
      - 1 to 4 words
      - All uppercase
      - No stop words
      - No word ending in -LY (adverb), -ED (past verb), -ING (gerund/present verb)
        unless it is a known valid suffix (e.g. KING ends in -ING but is valid — 
        we only reject if the word itself is 5+ chars ending in those suffixes)
    """
    words = candidate.strip().split()
    if not (1 <= len(words) <= 4):
        return False
    if candidate != candidate.upper():
        return False
    for w in words:
        if w in SPEAKER_STOP_WORDS:
            return False
        # Reject adverbs ending in -LY (e.g. SAVAGELY, WOOINGLY)
        if len(w) > 4 and w.endswith('LY'):
            return False
        # Reject past-tense verbs ending in -ED (e.g. ACCOUNTED, SLAUGHTER'D excluded by apostrophe check)
        if len(w) > 4 and w.endswith('ED'):
            return False
        # Reject gerunds/present verbs ending in -ING but allow KING, RING etc (len check)
        if len(w) > 5 and w.endswith('ING') and w not in ('NOTHING', 'SOMETHING'):
            return False
    return True

ROMAN_MAP = {
    "I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6, "VII": 7,
    "VIII": 8, "IX": 9, "X": 10, "XI": 11, "XII": 12,
    "i": 1, "ii": 2, "iii": 3, "iv": 4, "v": 5, "vi": 6, "vii": 7,
    "viii": 8, "ix": 9, "x": 10, "xi": 11, "xii": 12,
}


def roman_to_int(s):
    """Convert Roman numeral string to int. Returns None if invalid."""
    return ROMAN_MAP.get(s.strip().upper())


def parse_act_scene(key):
    """
    Parse scene key to (act_num, scene_num) for sorting.
    Handles: ACT 1 SCENE 1, ACT 1,SCENE 1, ACT 1, SCENE 1, ACT 2 ,SCENE 1, ACT 3, SCN 3, ACT IV, SCENE VII
    """
    # Match ACT <number/roman> ... SCENE or SCN <number/roman>
    m = re.search(r"ACT\s*[,]?\s*(\d+|IV?X?I{0,3})\s*[,]?\s*(?:SCENE|SCN)\s*[,]?\s*(\d+|IV?X?I{0,3})", key, re.I)
    if m:
        a, s = m.group(1), m.group(2)
        act = int(a) if a.isdigit() else roman_to_int(a)
        scene = int(s) if s.isdigit() else roman_to_int(s)
        if act is not None and scene is not None:
            return (act, scene)
    # Fallback: extract any two numbers
    nums = re.findall(r"\d+", key)
    if len(nums) >= 2:
        return (int(nums[0]), int(nums[1]))
    return (0, 0)


def format_scene_title(key):
    """Format scene key as 'ACT X, SCENE Y'."""
    m = re.search(r"ACT\s*[,]?\s*(\d+|IV?X?I{0,3})\s*[,]?\s*(?:SCENE|SCN)\s*[,]?\s*(\d+|IV?X?I{0,3})", key, re.I)
    if m:
        return f"ACT {m.group(1).upper()}, SCENE {m.group(2).upper()}"
    return key


def add_highlighted_heading(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    try:
        from docx.enum.text import WD_COLOR_INDEX
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    except Exception:
        pass
    return para


def add_section_heading(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    return para


def add_separator(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return para


def fix_allcaps_dialogue(text):
    """
    Convert erroneously ALL-CAPS dialogue to sentence case.
    Handles: "DOUBLY REDOUBLED STROKES UPON THE FOE:", "SHOW'D LIKE A REBEL'S WHORE: but all's too weak:"
    Skips stage directions [brackets] and already mixed-case text.
    """
    if not text or not text.strip():
        return text
    if text.strip().startswith("["):
        return text
    # If entire line is all caps, convert to sentence case
    letters = [c for c in text if c.isalpha()]
    if letters and all(c.isupper() for c in letters):
        if len(text) > 1:
            return text[0].upper() + text[1:].lower()
        return text
    # If line has colon: convert only the part BEFORE colon if that part is all caps
    if ":" in text:
        idx = text.index(":")
        before, after = text[:idx], text[idx:]
        before_letters = [c for c in before if c.isalpha()]
        if before_letters and all(c.isupper() for c in before_letters) and len(before.strip()) > 1:
            before = before[0].upper() + before[1:].lower() if before.strip() else before
            return before + after
    return text


def get_sortable_line_keys(scene_data):
    """Return list of (sort_key, line_key) for all lines. Handles numeric and non-numeric keys."""
    result = []
    for k in scene_data.keys():
        try:
            n = int(k)
            result.append((n, k))
        except ValueError:
            result.append((999999, k))  # Put non-numeric at end
    result.sort(key=lambda x: x[0])
    return [lk for _, lk in result]


def create_docx(json_path, title, output_path):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    add_highlighted_heading(doc, title)
    doc.add_paragraph()

    # DRAMATIS PERSONAE
    if "DRAMATIS PERSONAE" in data:
        add_section_heading(doc, "DRAMATIS PERSONAE")
        doc.add_paragraph()
        dp = data["DRAMATIS PERSONAE"]
        line_keys = get_sortable_line_keys(dp)
        for line_key in line_keys:
            text = dp[line_key].get("play", "")
            if text:
                doc.add_paragraph(text)
        doc.add_paragraph()

    # Scenes
    scene_keys = [k for k in data.keys() if k.startswith("ACT")]
    scene_keys.sort(key=parse_act_scene)

    for scene_key in scene_keys:
        scene_title = format_scene_title(scene_key)
        add_section_heading(doc, scene_title)
        doc.add_paragraph()
        add_separator(doc, "======PLAY TEXT======")
        doc.add_paragraph()

        scene_data = data[scene_key]
        line_keys = get_sortable_line_keys(scene_data)

        for line_key in line_keys:
            play_text = scene_data[line_key].get("play", "")
            para = doc.add_paragraph()
            run_num = para.add_run(f"{line_key}. ")
            run_num.font.size = Pt(11)
            # Detect real speaker names (e.g. "MACBETH: If it were done...")
            # Uses is_speaker_name() to avoid bolding verse fragments or scholar names
            speaker_match = re.match(r"^([A-Z][A-Z '\-\.]{0,35}):(.*)$", play_text or "")
            formatted = False
            if speaker_match:
                speaker = speaker_match.group(1).strip()
                rest = speaker_match.group(2)
                if is_speaker_name(speaker):
                    formatted = True
                    run_speaker = para.add_run(speaker + ": ")
                    run_speaker.bold = True
                    run_speaker.font.size = Pt(11)
                    run_rest = para.add_run(rest.lstrip())
                    run_rest.font.size = Pt(11)
            if not formatted:
                display = fix_allcaps_dialogue(play_text if play_text else "")
                run_text = para.add_run(display)
                run_text.font.size = Pt(11)

        doc.add_paragraph()
        add_separator(doc, "========SCHOLARLY COMMENTARY========")
        doc.add_paragraph()
        doc.add_paragraph()

    doc.save(output_path)
    return output_path


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    data_dir = os.path.join(script_dir, "Public", "Data")

    if len(sys.argv) >= 2:
        slug = sys.argv[1]
        title = sys.argv[2] if len(sys.argv) >= 3 else slug.replace("_", " ").title()
    else:
        slug = "twelfth_night"
        title = "Twelfth Night"

    json_name = f"{slug}.json"
    json_path = os.path.join(data_dir, json_name)

    if not os.path.isfile(json_path):
        print(f"JSON not found: {json_path}")
        sys.exit(1)

    safe_title = title.replace(",", "").replace(" ", "_")
    output_name = f"{safe_title}_With_Line_Numbers.docx"
    output_path = os.path.join(script_dir, output_name)

    print(f"Reading {json_path}...")
    print(f"Creating {output_path}...")
    create_docx(json_path, title, output_path)
    print(f"Done! Saved: {output_path}")


if __name__ == "__main__":
    main()
